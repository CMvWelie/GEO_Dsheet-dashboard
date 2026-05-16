"""Tests voor WordHoofdstukExporter."""
from __future__ import annotations
import sys, os, tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reporting.models import ReportSection, ReportField, ReportTable, ReportImageRequest, ReportMetadata, TextBlock
from exporters.word_hoofdstuk_exporter import WordHoofdstukExporter


TEMPLATE = os.path.join(os.path.dirname(__file__), '..', 'templates', 'damwand_stijlen.docx')


def _export(secties, metadata=None) -> Document:
    """Exporteer naar een tijdelijk bestand en lees terug."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    exp = WordHoofdstukExporter()
    fout = exp.export(
        sections=secties,
        metadata=metadata or ReportMetadata(project_name='Testproject'),
        project=None,
        template_path=TEMPLATE,
        output_path=pad,
    )
    assert fout is None, f'Export fout: {fout}'
    doc = Document(pad)
    os.unlink(pad)
    return doc


def test_export_maakt_bestand_aan() -> None:
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    exp = WordHoofdstukExporter()
    fout = exp.export(
        sections=[],
        metadata=ReportMetadata(project_name='P'),
        project=None,
        template_path=TEMPLATE,
        output_path=pad,
    )
    assert fout is None
    assert os.path.exists(pad)
    os.unlink(pad)


def test_sectietitel_wordt_heading() -> None:
    sec = ReportSection(id='test', title='Mijn Sectie')
    doc = _export([sec])
    teksten = [p.text for p in doc.paragraphs]
    assert 'Mijn Sectie' in teksten


def test_veld_wordt_paragraaf() -> None:
    sec = ReportSection(id='test', title='Sectie')
    sec.fields.append(ReportField('k', 'Profiel', 'AZ 14-700', ''))
    doc = _export([sec])
    tekst = ' '.join(p.text for p in doc.paragraphs)
    assert 'Profiel' in tekst
    assert 'AZ 14-700' in tekst


def test_tabel_kolommen_aanwezig() -> None:
    sec = ReportSection(id='test', title='Sectie')
    sec.tables.append(ReportTable(
        id='t', title='',
        columns=['Fase', 'Moment'],
        rows=[['F1', '100']],
    ))
    doc = _export([sec])
    assert len(doc.tables) >= 1
    header_cellen = [c.text for c in doc.tables[0].rows[0].cells]
    assert 'Fase' in header_cellen
    assert 'Moment' in header_cellen


def test_grondsoorten_v2_overzicht_lijnt_laagdata_links_uit() -> None:
    sec = ReportSection(id='grondsoorten_v2_overzicht', title='Grondsoorten')
    sec.tables.append(ReportTable(
        id='grondsoorten_v2_overzicht_tabel',
        title='',
        columns=['Laag', 'gamma'],
        rows=[['Zand', '18,0']],
    ))

    doc = _export([sec])

    assert doc.tables[0].rows[1].cells[0].paragraphs[0].alignment == (
        WD_ALIGN_PARAGRAPH.LEFT
    )


def test_grondsoorten_v2_overzicht_begint_op_nieuwe_pagina() -> None:
    sec = ReportSection(
        id='grondsoorten_v2_overzicht',
        title='Grondsoortentabel v2 - Grondsoorten',
    )

    doc = _export([ReportSection(id='vooraf', title='Vooraf'), sec])

    kop = next(
        p for p in doc.paragraphs
        if p.text == 'Grondsoortentabel v2 - Grondsoorten'
    )
    assert kop.paragraph_format.page_break_before is True


def test_tabel_met_kolomgroepen_wordt_geschreven() -> None:
    from docx.oxml.ns import qn

    sec = ReportSection(id='grondsoorten_v2_fase_1', title='Grondlaagopbouw fases')
    sec.tables.append(ReportTable(
        id='grondsoorten_v2_fase_1_tabel',
        title='',
        columns=['Laag', 'b.k. laag', 'o.k. laag', 'Laag', 'b.k. laag', 'o.k. laag'],
        rows=[['Zand', '0,00', 'Max', 'Klei', '0,00', 'Max']],
        column_groups=[
            ('Grondlagen linkerzijde', 3),
            ('Grondlagen rechterzijde', 3),
        ],
    ))

    doc = _export([sec])

    assert 'Grondlagen linkerzijde' in doc.tables[0].rows[0].cells[0].text
    assert 'Grondlagen rechterzijde' in doc.tables[0].rows[0].cells[3].text
    assert [col.get(qn('w:w')) for col in doc.tables[0]._tbl.tblGrid.gridCol_lst] == [
        '2268', '1134', '1134', '2268', '1134', '1134',
    ]
    assert doc.tables[0].rows[0].cells[0]._tc.tcPr.tcW.w == 4536
    assert doc.tables[0].rows[0].cells[3]._tc.tcPr.tcW.w == 4536
    assert [
        cell._tc.tcPr.tcW.w
        for cell in doc.tables[0].rows[2].cells
    ] == [2268, 1134, 1134, 2268, 1134, 1134]
    assert doc.tables[0].rows[2].cells[0].paragraphs[0].alignment == (
        WD_ALIGN_PARAGRAPH.LEFT
    )
    assert doc.tables[0].rows[2].cells[3].paragraphs[0].alignment == (
        WD_ALIGN_PARAGRAPH.LEFT
    )


def test_grondsoorten_v2_gelijke_zijden_krijgt_drie_wordkolommen() -> None:
    from docx.oxml.ns import qn

    sec = ReportSection(id='grondsoorten_v2_fase_1', title='Grondlaagopbouw fases')
    sec.tables.append(ReportTable(
        id='grondsoorten_v2_fase_1_tabel',
        title='',
        columns=['Laag', 'b.k. laag', 'o.k. laag'],
        rows=[['Zand', '0,00', 'Max']],
        column_groups=[('Grondlagen', 3)],
    ))

    doc = _export([sec])

    assert 'Grondlagen' in doc.tables[0].rows[0].cells[0].text
    assert [col.get(qn('w:w')) for col in doc.tables[0]._tbl.tblGrid.gridCol_lst] == [
        '2268', '1134', '1134',
    ]
    assert doc.tables[0].rows[0].cells[0]._tc.tcPr.tcW.w == 4536
    assert [
        cell._tc.tcPr.tcW.w
        for cell in doc.tables[0].rows[2].cells
    ] == [2268, 1134, 1134]
    assert doc.tables[0].rows[2].cells[0].paragraphs[0].alignment == (
        WD_ALIGN_PARAGRAPH.LEFT
    )


def test_grondsoorten_v2_word_cellen_kunnen_worden_doorgestreept() -> None:
    sec = ReportSection(id='grondsoorten_v2_fase_1', title='Grondlaagopbouw fases')
    sec.tables.append(ReportTable(
        id='grondsoorten_v2_fase_1_tabel',
        title='',
        columns=['Laag', 'b.k. laag', 'o.k. laag'],
        rows=[
            ['Zand', '0,00', '-5,00'],
            ['Klei', '-4,00', 'Max'],
        ],
        column_groups=[('Grondlagen', 3)],
        strikethrough_cells=[
            [True, True, True],
            [False, False, False],
        ],
    ))

    doc = _export([sec])

    assert doc.tables[0].rows[2].cells[0].paragraphs[0].runs[0].font.strike is True
    assert doc.tables[0].rows[2].cells[1].paragraphs[0].runs[0].font.strike is True
    assert doc.tables[0].rows[3].cells[0].paragraphs[0].runs[0].font.strike is not True


def test_grondsoorten_v2_ongewijzigde_zijde_wordt_samengevoegd() -> None:
    from docx.oxml.ns import qn

    sec = ReportSection(id='grondsoorten_v2_fase_2', title='')
    sec.tables.append(ReportTable(
        id='grondsoorten_v2_fase_2_tabel',
        title='',
        columns=['Laag', 'b.k. laag', 'o.k. laag', 'Laag', 'b.k. laag', 'o.k. laag'],
        rows=[
            [
                'Zand', '0,00', '-5,00',
                'Grondopbouw ongewijzigd t.o.v. vorige fase', '', '',
            ],
            ['Klei', '-5,00', 'Max', '', '', ''],
        ],
        column_groups=[
            ('Grondlagen linkerzijde', 3),
            ('Grondlagen rechterzijde', 3),
        ],
    ))

    doc = _export([sec])
    samengevoegd = doc.tables[0].rows[2].cells[3]
    tc_pr = samengevoegd._tc.tcPr

    assert samengevoegd.text == 'Grondopbouw ongewijzigd t.o.v. vorige fase'
    assert tc_pr.gridSpan.val == 3
    assert tc_pr.tcW.w == 4536
    assert tc_pr.find(qn('w:vMerge')).get(qn('w:val')) == 'restart'
    assert tc_pr.find(qn('w:vAlign')).get(qn('w:val')) == 'center'


def test_grondsoorten_v2_fase_intro_voor_tabel_met_word_bullets() -> None:
    from docx.oxml.ns import qn

    sec = ReportSection(id='grondsoorten_v2_fase_1', title='Grondlaagopbouw fases')
    sec.text_blocks.append(TextBlock(
        id='intro',
        section='grondsoorten_v2_fase_1',
        generated_text=(
            'Het volgende profiel wordt gehanteerd in de volgende fases:\n'
            'Fase 1\n'
            'Fase 2'
        ),
    ))
    sec.tables.append(ReportTable(
        id='t',
        title='',
        columns=['Laag', 'b.k. laag', 'o.k. laag'],
        rows=[['Zand', '0,00', 'Max']],
    ))

    doc = _export([sec])
    teksten = [p.text for p in doc.paragraphs if p.text.strip()]

    assert teksten[:4] == [
        'Testproject',
        'Grondlaagopbouw fases',
        'Het volgende profiel wordt gehanteerd in de volgende fases:',
        'Fase 1',
    ]
    assert teksten[4] == 'Fase 2'
    for tekst in ('Fase 1', 'Fase 2'):
        para = next(p for p in doc.paragraphs if p.text == tekst)
        num_pr = para._element.pPr.find(qn('w:numPr'))
        assert num_pr is not None
    alle_teksten = [p.text for p in doc.paragraphs]
    assert alle_teksten[alle_teksten.index('Fase 2') + 1] == ''
    assert any(
        lvl_text.get(qn('w:val')) == '-'
        for lvl_text in doc.part.numbering_part.element.iter(qn('w:lvlText'))
    )


def test_grondsoorten_v2_tweede_fasetabel_krijgt_geen_extra_heading() -> None:
    eerste = ReportSection(id='grondsoorten_v2_fase_1', title='Grondlaagopbouw fases')
    eerste.text_blocks.append(TextBlock(
        id='intro_1',
        section='grondsoorten_v2_fase_1',
        generated_text='In de fase "Fase 1" wordt het volgende profiel gehanteerd:',
    ))
    eerste.tables.append(ReportTable(
        id='grondsoorten_v2_fase_1_tabel',
        title='',
        columns=['Laag', 'b.k. laag', 'o.k. laag'],
        rows=[['Zand', '0,00', 'Max']],
    ))
    tweede = ReportSection(id='grondsoorten_v2_fase_2', title='')
    tweede.text_blocks.append(TextBlock(
        id='intro_2',
        section='grondsoorten_v2_fase_2',
        generated_text='In de fase "Fase 2" wordt het volgende profiel gehanteerd:',
    ))
    tweede.tables.append(ReportTable(
        id='grondsoorten_v2_fase_2_tabel',
        title='',
        columns=['Laag', 'b.k. laag', 'o.k. laag'],
        rows=[['Klei', '-1,00', 'Max']],
    ))

    doc = _export([eerste, tweede])
    teksten = [p.text for p in doc.paragraphs]

    assert teksten.count('Grondlaagopbouw fases') == 1
    assert 'In de fase "Fase 1" wordt het volgende profiel gehanteerd:' in teksten
    assert 'In de fase "Fase 2" wordt het volgende profiel gehanteerd:' in teksten
    tweede_intro_index = teksten.index(
        'In de fase "Fase 2" wordt het volgende profiel gehanteerd:'
    )
    assert teksten[tweede_intro_index - 1] == ''
    assert teksten[tweede_intro_index + 1] == ''
    assert len(doc.tables) == 2


def test_export_zonder_template_gebruikt_lege_doc() -> None:
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    fout = WordHoofdstukExporter().export(
        sections=[ReportSection(id='x', title='Y')],
        metadata=ReportMetadata(),
        project=None,
        template_path=None,
        output_path=pad,
    )
    assert fout is None
    os.unlink(pad)


def test_png_hoogte_cm_berekent_correct() -> None:
    import struct
    import zlib

    from exporters.word_hoofdstuk_exporter import _png_hoogte_cm

    def _mini_png(w: int, h: int) -> bytes:
        header = b'\x89PNG\r\n\x1a\n'
        ihdr_data = struct.pack('>IIBBBBB', w, h, 8, 2, 0, 0, 0)
        crc = zlib.crc32(b'IHDR' + ihdr_data) & 0xFFFFFFFF
        ihdr = struct.pack('>I', 13) + b'IHDR' + ihdr_data + struct.pack('>I', crc)
        return header + ihdr

    png = _mini_png(100, 200)
    result = _png_hoogte_cm(png, 6.0)
    assert abs(result - 12.0) < 0.001


def test_fase_invoer_sectie_maakt_tabel_in_word() -> None:
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    kaart = FaseCard(fase_num=1, stage_name='Fase 1: Test')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    kaart.rows.append(FaseRow('Leganker', '-0,3 [m NAP]', '0 graden t.o.v. maaiveld'))
    sec = FaseInvoerSectie(id='fase_1', title='Fase 1: Test', fase_card=kaart)

    doc = _export([sec], metadata=ReportMetadata(project_name='Test'))
    assert len(doc.tables) >= 1


def test_fase_invoer_secties_krijgen_gezamenlijke_fasering_intro() -> None:
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    kaart_1 = FaseCard(fase_num=1, stage_name='Initieel')
    kaart_1.rows.append(FaseRow('Maaiveld Links', '0,0 [m NAP]'))
    kaart_2 = FaseCard(fase_num=2, stage_name='Maatgevende fase')
    kaart_2.rows.append(FaseRow('Water Rechts', '-1,0 [m NAP]'))
    secties = [
        FaseInvoerSectie(id='fase_1', title='Fase 1', fase_card=kaart_1),
        FaseInvoerSectie(id='fase_2', title='Fase 2', fase_card=kaart_2),
    ]

    doc = _export(secties, metadata=ReportMetadata(project_name='Test'))
    teksten = [p.text for p in doc.paragraphs if p.text.strip()]

    assert 'Fasering' in teksten
    assert 'De onderstaande fasering is als volgt toegepast.' in teksten
    assert 'Fase 1: Initieel' in teksten
    assert 'Fase 2: Maatgevende fase' in teksten
    assert 'Voor de grafisch weergegeven bouwfasering(en) zie de onderstaande tabel.' in teksten
    assert 'Fase 1' not in teksten
    assert len(doc.tables) == 2


def test_fasering_intro_noemt_alle_projectfases_ook_bij_deelselectie() -> None:
    from parsers.models import FileBundle, Project, ResultSummary, Stage
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    secties = []
    for nr, naam in enumerate(['Initieel', 'Ontgraven'], start=1):
        kaart = FaseCard(fase_num=nr, stage_name=naam)
        kaart.rows.append(FaseRow('Maaiveld Links', '0,0 [m NAP]'))
        secties.append(FaseInvoerSectie(
            id=f'fase_{nr}',
            title=f'Fase {nr}',
            fase_card=kaart,
        ))

    project = Project(
        base_name='t',
        project_name='T',
        file_bundle=FileBundle(),
        stages=[
            Stage(name='Initieel'),
            Stage(name='Ontgraven'),
            Stage(name='Stempel plaatsen'),
            Stage(name='Eindsituatie'),
        ],
        result_summaries=[
            ResultSummary(
                stage_number=5,
                max_moment_knm=0.0,
                max_shear_kn=0.0,
                max_disp_mm=0.0,
                mob_moment_pct=0.0,
                mob_grond_pct=0.0,
            )
        ],
    )

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    fout = WordHoofdstukExporter().export(
        sections=secties,
        metadata=ReportMetadata(project_name='Test'),
        project=project,
        template_path=None,
        output_path=pad,
    )
    assert fout is None
    doc = Document(pad)
    os.unlink(pad)
    teksten = [p.text for p in doc.paragraphs if p.text.strip()]

    assert 'Fase 1: Initieel' in teksten
    assert 'Fase 2: Ontgraven' in teksten
    assert 'Fase 3: Stempel plaatsen' in teksten
    assert 'Fase 4: Eindsituatie' in teksten
    assert 'Fase 5: Fase 5' in teksten
    assert len(doc.tables) == 2


def test_fase_invoer_secties_houden_witregel_en_tabel_bij_elkaar() -> None:
    from docx.oxml.ns import qn
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    secties = []
    for nr, naam in enumerate(['Initieel', 'Maatgevende fase'], start=1):
        kaart = FaseCard(fase_num=nr, stage_name=naam)
        kaart.rows.append(FaseRow('Maaiveld Links', '0,0 [m NAP]'))
        secties.append(FaseInvoerSectie(
            id=f'fase_{nr}',
            title=f'Fase {nr}',
            fase_card=kaart,
        ))

    doc = _export(secties, metadata=ReportMetadata(project_name='Test'))

    assert any(
        p.text == '' and p.paragraph_format.keep_with_next
        for p in doc.paragraphs
    )
    eerste_paragraaf = doc.tables[0].rows[0].cells[0].paragraphs[0]
    keep_together = eerste_paragraaf._p.pPr.find(qn('w:keepLines'))
    keep_with_next = eerste_paragraaf._p.pPr.find(qn('w:keepNext'))
    assert keep_together is not None
    assert keep_with_next is not None


def test_damwandgegevens_sectie_wordt_voorbeeldtabel() -> None:
    from docx.oxml.ns import qn

    sec = ReportSection(id='damwand_gegevens', title='Damwandgegevens')
    sec.fields.extend([
        ReportField('profiel', 'Profiel', 'AZ 18-700'),
        ReportField('staalkwaliteit', 'Staalkwaliteit', 'S355'),
        ReportField('hoogte', 'Hoogte', '420,0', 'mm'),
        ReportField('ondersteuning_3', 'Anker A', '-2,1', 'm NAP'),
    ])

    doc = _export([sec], metadata=ReportMetadata(project_name='Test'))
    tbl = doc.tables[0]

    assert [col.get(qn('w:w')) for col in tbl._tbl.tblGrid.gridCol_lst] == [
        '2835', '1701', '1134',
    ]
    assert [cell.text for cell in tbl.rows[0].cells] == [
        'Parameter', 'Waarde', 'Eenheid',
    ]
    assert [cell.text for cell in tbl.rows[1].cells] == ['Profiel', 'AZ 18-700', '']
    assert [cell.text for cell in tbl.rows[3].cells] == ['Hoogte', '420,0', '[mm]']
    assert [cell.text for cell in tbl.rows[4].cells] == ['', '', '']
    assert [cell.text for cell in tbl.rows[5].cells] == [
        'Anker A', '-2,1', '[m NAP]',
    ]
    for cell in tbl.rows[0].cells:
        shd = cell._tc.tcPr.find(qn('w:shd'))
        assert shd is not None
        assert shd.get(qn('w:fill')) == '147ACF'
    assert tbl.rows[0]._tr.trPr.trHeight.get(qn('w:val')) == '255'
    assert tbl.rows[0]._tr.trPr.trHeight.get(qn('w:hRule')) == 'exact'
    assert tbl.rows[4]._tr.trPr.trHeight.get(qn('w:val')) == '40'


def test_damwandgegevens_sectie_schrijft_toelichting_rond_tabel() -> None:
    sec = ReportSection(id='damwand_gegevens', title='Damwandgegevens')
    sec.text_blocks.append(TextBlock(
        id='intro',
        section='damwand_gegevens',
        generated_text='Voor de grondkering zijn de volgende eigenschappen aangehouden.',
    ))
    sec.text_blocks.append(TextBlock(
        id='toelichting',
        section='damwand_gegevens',
        generated_text='Hierin is:\nEI\tde ongereduceerde buigstijfheid',
    ))
    sec.fields.append(ReportField('profiel', 'Profiel', 'AZ 18-700', '-'))

    doc = _export([sec], metadata=ReportMetadata(project_name='Test'))
    teksten = [p.text for p in doc.paragraphs]

    assert 'Voor de grondkering zijn de volgende eigenschappen aangehouden.' in teksten
    assert 'Hierin is:' in teksten
    assert 'EI\tde ongereduceerde buigstijfheid' in teksten
    assert doc.tables[0].rows[1].cells[2].text == '[-]'


def test_fase_sectie_tabel_bevat_kolomhoofden() -> None:
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    kaart = FaseCard(fase_num=2, stage_name='Fase 2: Belasting')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    kaart.rows.append(FaseRow(
        'Bovenbelasting',
        'op maaiveld',
        '5,0 [kN/m²]',
        extra_lines=['3,0m breed', '0m vanaf damwand'],
    ))
    sec = FaseInvoerSectie(id='fase_2', title='Fase 2: Belasting', fase_card=kaart)

    doc = _export([sec], metadata=ReportMetadata())
    tbl = doc.tables[0]
    teksten_rij1 = [c.text for c in tbl.rows[1].cells]
    assert 'Parameter' in teksten_rij1
    assert 'Niveau' in teksten_rij1
    assert 'Toelichting' in teksten_rij1
    assert len(tbl.rows) >= 5


def test_fase_sectie_tabel_gebruikt_voorbeeld_breedtes_en_headerkleur() -> None:
    from docx.oxml.ns import qn
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    kaart = FaseCard(fase_num=2, stage_name='Fase 2: Belasting')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    kaart.rows.append(FaseRow(
        'Bovenbelasting',
        'op maaiveld',
        '5,0 [kN/m²]',
        extra_lines=['3,0m breed', '0,0m vanaf damwand'],
    ))
    sec = FaseInvoerSectie(id='fase_2', title='Fase 2: Belasting', fase_card=kaart)

    doc = _export([sec], metadata=ReportMetadata())
    tbl = doc.tables[0]
    grid = tbl._tbl.tblGrid
    assert [col.get(qn('w:w')) for col in grid.gridCol_lst] == [
        '1701', '1134', '2835', '3572',
    ]
    assert tbl.rows[0].cells[0]._tc.tcPr.tcW.w == 5670
    assert tbl.rows[0].cells[3]._tc.tcPr.tcW.w == 3572
    assert tbl.rows[1].cells[0]._tc.tcPr.tcW.w == 1701
    assert tbl.rows[1].cells[1]._tc.tcPr.tcW.w == 1134
    assert tbl.rows[1].cells[2]._tc.tcPr.tcW.w == 2835

    for row_idx in (0, 1):
        for cell in tbl.rows[row_idx].cells:
            shd = cell._tc.tcPr.find(qn('w:shd'))
            assert shd is not None
            assert shd.get(qn('w:fill')) == '147ACF'

    header_run = tbl.rows[1].cells[0].paragraphs[0].runs[0]
    assert header_run.font.size.pt == 8
    data_run = tbl.rows[2].cells[0].paragraphs[0].runs[0]
    assert data_run.font.size.pt == 7

    heights = [
        tbl.rows[i]._tr.trPr.trHeight.get(qn('w:val'))
        for i in range(len(tbl.rows))
        if tbl.rows[i]._tr.trPr is not None
        and tbl.rows[i]._tr.trPr.trHeight is not None
    ]
    assert all(h == '255' for h in heights)


def test_fase_sectie_gebruikt_theme_fontgroottes(monkeypatch) -> None:
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie
    from ui import table_styles

    monkeypatch.setattr(table_styles, 'WORD_BODY_TEXT_SIZE', 12)
    monkeypatch.setattr(table_styles, 'WORD_TABLE_TEXT_SIZE', 6)
    monkeypatch.setattr(table_styles, 'WORD_TABLE_HEADER_SIZE', 9)

    kaart = FaseCard(fase_num=1, stage_name='Fase 1')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    sec = FaseInvoerSectie(id='fase_1', title='Fase 1', fase_card=kaart)

    doc = _export([sec], metadata=ReportMetadata())
    tbl = doc.tables[0]

    assert doc.styles['Normal'].font.size.pt == 12
    assert tbl.rows[1].cells[0].paragraphs[0].runs[0].font.size.pt == 9
    assert tbl.rows[2].cells[0].paragraphs[0].runs[0].font.size.pt == 6


def test_resultaatbeschrijving_specificaties_tabel_in_word() -> None:
    from docx.oxml.ns import qn
    from parsers.models import (
        FileBundle,
        Project,
        ResultPoint,
        ResultStage,
        ResultStep,
        ResultSummary,
        SheetPilingElement,
    )

    project = Project(
        base_name='t',
        project_name='T',
        file_bundle=FileBundle(),
        sheet_piling=[
            SheetPilingElement(
                name='AZ 18 (S240GP)',
                x=0.0,
                bottom=-12.0,
                top=1.0,
                width=1.0,
                opneembaar_moment_knm=250.0,
                steel_quality='S240GP',
            )
        ],
        result_summaries=[
            ResultSummary(
                stage_number=1,
                max_moment_knm=100.0,
                max_shear_kn=50.0,
                max_disp_mm=10.0,
                mob_moment_pct=60.0,
                mob_grond_pct=70.0,
            )
        ],
        result_steps={
            'CUR 166 6.4': ResultStep(
                raw_step='CUR 166 6.4',
                stages={
                    1: ResultStage(
                        stage_number=1,
                        points=[ResultPoint(depth=-2.0, moment=-123.0, shear=45.0, disp=4.0)],
                    )
                },
            ),
            'CUR 166 6.5': ResultStep(
                raw_step='CUR 166 6.5',
                stages={
                    1: ResultStage(
                        stage_number=1,
                        points=[ResultPoint(depth=-1.0, moment=10.0, shear=5.0, disp=-9.0)],
                    )
                },
            ),
        },
    )
    sec = ReportSection(id='conclusietabel', title='Resultaten per fase')

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    fout = WordHoofdstukExporter().export(
        sections=[sec],
        metadata=ReportMetadata(project_name='Test'),
        project=project,
        template_path=None,
        output_path=pad,
    )
    assert fout is None
    doc = Document(pad)
    os.unlink(pad)

    tbl = doc.tables[0]
    assert [col.get(qn('w:w')) for col in tbl._tbl.tblGrid.gridCol_lst] == [
        '2835', '1418', '1984', '1134',
    ]
    assert tbl.rows[0].cells[0].text == 'Grondkering'
    assert tbl.rows[0].cells[0]._tc.tcPr.tcW.w == 7371   # merged over alle 4 kolommen
    assert tbl.rows[0].cells[0]._tc.tcPr.find(qn('w:shd')).get(qn('w:fill')) == '147ACF'
    # kolom 0+1 samengevoegd → python-docx herhaalt tekst bij beide indices
    assert [cell.text for cell in tbl.rows[1].cells] == ['Profiel', 'Profiel', 'AZ 18', '[-]']
    assert tbl.rows[7].cells[0].text == 'Resultaten'
    assert tbl.rows[7].cells[0]._tc.tcPr.find(qn('w:shd')).get(qn('w:fill')) == '147ACF'
    assert tbl.rows[7].cells[1].text == 'Verificatiestap'  # stap-header in kolom 1
    assert [cell.text for cell in tbl.rows[8].cells] == [
        'Moment Msd UGT', 'stap 6.4', '123,0', '[kNm/m]',
    ]
    assert [cell.text for cell in tbl.rows[10].cells] == [
        'Verplaatsing urep BGT', 'stap 6.5', '9,0', '[mm]',
    ]
    assert [cell.text for cell in tbl.rows[11].cells] == [
        'Gemobiliseerd Moment', 'Gemobiliseerd Moment', '60,0', '[%]',
    ]


def test_fase_sectie_voegt_geen_paddingrij_toe_zonder_afbeelding() -> None:
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    kaart = FaseCard(fase_num=1, stage_name='Fase 1')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    sec = FaseInvoerSectie(id='fase_1', title='Fase 1', fase_card=kaart)

    doc = _export([sec], metadata=ReportMetadata())
    tbl = doc.tables[0]
    assert len(tbl.rows) == 3  # twee koprijen + één gevulde tekstregel


def test_fase_sectie_voegt_paddingrij_toe_als_afbeelding_hoger_is(monkeypatch) -> None:
    import base64
    from exporters import word_hoofdstuk_exporter
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    png_1x1 = base64.b64decode(
        'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwAD'
        'hgGAWjR9awAAAABJRU5ErkJggg=='
    )

    monkeypatch.setattr(
        word_hoofdstuk_exporter,
        'render_figuur',
        lambda _img, _project: png_1x1,
    )
    monkeypatch.setattr(word_hoofdstuk_exporter, '_png_hoogte_cm', lambda *_: 12.0)
    kaart = FaseCard(fase_num=1, stage_name='Fase 1')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    sec = FaseInvoerSectie(id='fase_1', title='Fase 1', fase_card=kaart)

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    fout = WordHoofdstukExporter().export(
        sections=[sec],
        metadata=ReportMetadata(),
        project=object(),
        template_path=None,
        output_path=pad,
    )
    assert fout is None
    doc = Document(pad)
    os.unlink(pad)
    tbl = doc.tables[0]
    assert len(tbl.rows) == 4  # twee koprijen + één tekstregel + één paddingrij


def test_fase_sectie_voegt_geen_paddingrij_toe_als_afbeelding_lager_is(monkeypatch) -> None:
    import base64
    from exporters import word_hoofdstuk_exporter
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    png_1x1 = base64.b64decode(
        'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwAD'
        'hgGAWjR9awAAAABJRU5ErkJggg=='
    )
    monkeypatch.setattr(
        word_hoofdstuk_exporter,
        'render_figuur',
        lambda _img, _project: png_1x1,
    )
    monkeypatch.setattr(word_hoofdstuk_exporter, '_png_hoogte_cm', lambda *_: 0.1)

    kaart = FaseCard(fase_num=1, stage_name='Fase 1')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    sec = FaseInvoerSectie(id='fase_1', title='Fase 1', fase_card=kaart)

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    fout = WordHoofdstukExporter().export(
        sections=[sec],
        metadata=ReportMetadata(),
        project=object(),
        template_path=None,
        output_path=pad,
    )
    assert fout is None
    doc = Document(pad)
    os.unlink(pad)
    tbl = doc.tables[0]
    assert len(tbl.rows) == 3  # twee koprijen + één tekstregel, geen paddingrij


# ---------------------------------------------------------------------------
# Taak 6: figuurrendering
# ---------------------------------------------------------------------------

from parsers.models import Project, FileBundle, SheetPilingElement, Stage, Surface


def _mini_project() -> Project:
    return Project(
        base_name='t', project_name='T', file_bundle=FileBundle(),
        sheet_piling=[SheetPilingElement(
            name='AZ 14-700', x=0.0, bottom=-10.0, top=-2.0, width=1.4,
        )],
        stages=[Stage(name='F1')],
        surfaces=[Surface(nr=1, name='MV', points=[{'x': -10, 'y': 0}, {'x': 10, 'y': 0}])],
    )


def test_figuur_placeholder_zonder_project() -> None:
    sec = ReportSection(id='test', title='Grafieken')
    sec.images.append(ReportImageRequest(
        id='fig1', caption='Dwarsdoorsnede fase 1',
        figure_key='section', stage_index=0, step_key=None,
    ))
    doc = _export([sec], metadata=ReportMetadata())
    tekst = ' '.join(p.text for p in doc.paragraphs)
    assert 'Figuur' in tekst


def test_render_figuur_section_geeft_bytes() -> None:
    from reporting.models import ReportImageRequest as RIR
    from exporters.word_hoofdstuk_exporter import WordHoofdstukExporter
    project = _mini_project()
    img_req = RIR(id='f', caption='c', figure_key='section', stage_index=0, step_key=None)
    exp = WordHoofdstukExporter()
    result = exp._render_figuur(img_req, project)
    assert result is not None
    assert result[:4] == b'\x89PNG'


def test_render_figuur_moment_shear_geeft_bytes() -> None:
    from reporting.models import ReportImageRequest as RIR
    from exporters.word_hoofdstuk_exporter import WordHoofdstukExporter
    project = _mini_project()
    img_req = RIR(id='f', caption='c', figure_key='moment_shear', stage_index=0, step_key=None)
    exp = WordHoofdstukExporter()
    result = exp._render_figuur(img_req, project)
    assert result is not None
    assert result[:4] == b'\x89PNG'


def test_grondsoorten_v1_unit_groups_worden_geschreven() -> None:
    sec = ReportSection(id='soil_table_links', title='Grondsoortentabel — Links')
    sec.tables.append(ReportTable(
        id='soil_table_links_tabel',
        title='',
        columns=['Laag', 'BK laag', 'OK laag', 'Γd / yn', "c'kar",
                 "φ'kar", 'δ', 'kh1', 'kh2', 'kh3'],
        rows=[['Zand', '0,00', '-5,00', '15,0 / 18,0', '1,0', '32,5', '16,0', '10000', '5000', '2000']],
        unit_groups=[
            ('', 1),
            ('[m NAP]', 2),
            ('[kN/m³]', 1),
            ('[kN/m²]', 1),
            ('[°]', 2),
            ('[kN/m³]', 3),
        ],
    ))

    doc = _export([sec])

    tbl = doc.tables[0]
    # Rij 0: kolomnamen
    assert tbl.rows[0].cells[0].text == 'Laag'
    assert tbl.rows[0].cells[1].text == 'BK laag'
    # Rij 1: eenheden
    assert tbl.rows[1].cells[1].text == '[m NAP]'   # BK+OK samengevoegd
    assert tbl.rows[1].cells[5].text == '[°]'        # φ'+δ samengevoegd
    assert tbl.rows[1].cells[7].text == '[kN/m³]'   # kh1+2+3 samengevoegd
    # Rij 2: data
    assert tbl.rows[2].cells[0].text == 'Zand'
    assert tbl.rows[2].cells[3].text == '15,0 / 18,0'


def test_grondsoorten_v1_laag_kolom_links_uitgelijnd() -> None:
    sec = ReportSection(id='soil_table_links', title='Grondsoortentabel — Links')
    sec.tables.append(ReportTable(
        id='soil_table_links_tabel',
        title='',
        columns=['Laag', 'BK laag', 'OK laag', 'Γd / yn', "c'kar",
                 "φ'kar", 'δ', 'kh1', 'kh2', 'kh3'],
        rows=[['Zand', '0,00', '-5,00', '15,0 / 18,0', '1,0', '32,5', '16,0', '10000', '5000', '2000']],
        unit_groups=[
            ('', 1), ('[m NAP]', 2), ('[kN/m³]', 1),
            ('[kN/m²]', 1), ('[°]', 2), ('[kN/m³]', 3),
        ],
    ))

    doc = _export([sec])

    # Laag-kolom (col 0) in datarij moet links zijn
    assert doc.tables[0].rows[2].cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
