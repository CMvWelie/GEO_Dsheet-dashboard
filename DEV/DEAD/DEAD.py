# DEAD.py
"""Archief van dode code verwijderd tijdens opruiming op 2026-05-07.

Dit bestand is alleen bedoeld als terugvalarchief. Importeren of uitvoeren is niet nodig.
Elke entry bevat de originele broninhoud van een verwijderd of gesnoeid bestand.
"""

DEAD_CODE_ARCHIVE = {
    'ui/sidebar.py': r'''
"""Linker zijbalk: bestandsimport, projectselectie en fase-selectie."""

from __future__ import annotations
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QComboBox,
    QPushButton, QGroupBox
)
from PyQt6.QtCore import pyqtSignal, Qt

from ui.file_list_widget import FileListWidget


class StatusLabel(QLabel):
    """Gekleurde statusbadge (OK / WARN / ERR)."""

    _COLORS = {
        'ok': ('#2f7d32', '#eef8ef', '#9ac89c'),
        'warn': ('#b26a00', '#fff7eb', '#e2be83'),
        'err': ('#b42318', '#fff2f1', '#e5a6a1'),
        'idle': ('#555555', '#f4f4f4', '#cccccc'),
    }

    def set_status(self, status_type: str, text: str, detail: str = '') -> None:
        """Stel de statustekst en kleur in.

        Parameters
        ----------
        status_type: 'ok', 'warn', 'err' of 'idle'.
        text:        Korte statustekst.
        detail:      Optionele detailtekst als tooltip.
        """
        fg, bg, border = self._COLORS.get(status_type, self._COLORS['idle'])
        self.setStyleSheet(
            f'color: {fg}; background: {bg}; border: 1px solid {border};'
            f'border-radius: 8px; padding: 4px 8px; font-weight: bold; font-size: 12px;'
        )
        self.setText(text)
        if detail:
            self.setToolTip(detail)


class Sidebar(QWidget):
    """Linker paneel met bestandsimport, project- en fase-selectie.

    Signals
    -------
    import_requested:    Gebruiker wil bestanden importeren.
    process_requested:   Gebruiker wil bestanden verwerken.
    reset_requested:     Gebruiker wil alles resetten.
    project_changed:     Ander project geselecteerd (waarde = base_name).
    stage_changed:       Andere fase geselecteerd (waarde = index).
    """

    import_requested = pyqtSignal()
    process_requested = pyqtSignal()
    reset_requested = pyqtSignal()
    project_changed = pyqtSignal(str)
    stage_changed = pyqtSignal(int)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedWidth(300)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(8, 8, 8, 8)

        # Titel
        title = QLabel('<b>D-Sheet Dashboard</b>')
        title.setStyleSheet('font-size: 16px;')
        layout.addWidget(title)

        # Bestandsimport
        import_box = QGroupBox('Bestanden')
        import_layout = QVBoxLayout(import_box)

        btn_row = QWidget()
        btn_row_l = QHBoxLayout(btn_row)
        btn_row_l.setContentsMargins(0, 0, 0, 0)
        self._import_btn = QPushButton('Importeerâ€¦')
        self._import_btn.clicked.connect(self.import_requested.emit)
        self._process_btn = QPushButton('Verwerk')
        self._process_btn.clicked.connect(self.process_requested.emit)
        self._reset_btn = QPushButton('Reset')
        self._reset_btn.clicked.connect(self.reset_requested.emit)
        btn_row_l.addWidget(self._import_btn)
        btn_row_l.addWidget(self._process_btn)
        btn_row_l.addWidget(self._reset_btn)
        import_layout.addWidget(btn_row)

        self._file_list = FileListWidget()
        import_layout.addWidget(self._file_list)
        layout.addWidget(import_box)

        # Project selectie
        proj_box = QGroupBox('Project')
        proj_layout = QVBoxLayout(proj_box)
        self._project_combo = QComboBox()
        self._project_combo.currentTextChanged.connect(self._on_project_changed)
        proj_layout.addWidget(self._project_combo)
        layout.addWidget(proj_box)

        # Fase selectie
        stage_box = QGroupBox('Bouwfase')
        stage_layout = QVBoxLayout(stage_box)
        self._stage_combo = QComboBox()
        self._stage_combo.currentIndexChanged.connect(self._on_stage_changed)
        stage_layout.addWidget(self._stage_combo)
        layout.addWidget(stage_box)

        # Status
        self._status = StatusLabel('Gereed')
        self._status.set_status('idle', 'Gereed', 'Importeer bestanden om te beginnen.')
        layout.addWidget(self._status)

        layout.addStretch()

    def _on_project_changed(self, text: str) -> None:
        if text:
            self.project_changed.emit(text)

    def _on_stage_changed(self, index: int) -> None:
        self.stage_changed.emit(index)

    def set_status(self, status_type: str, text: str, detail: str = '') -> None:
        """Delegeer naar StatusLabel."""
        self._status.set_status(status_type, text, detail)

    def set_files(self, files: dict[str, str]) -> None:
        """Vul de bestandslijst.

        Parameters
        ----------
        files: Dict filename â†’ raw text.
        """
        self._file_list.set_files(files)

    def set_projects(self, project_names: dict[str, str]) -> None:
        """Vul de projectcombobox.

        Parameters
        ----------
        project_names: Dict base_name â†’ projectnaam.
        """
        self._project_combo.blockSignals(True)
        self._project_combo.clear()
        for base_name, display_name in project_names.items():
            self._project_combo.addItem(display_name, userData=base_name)
        self._project_combo.blockSignals(False)
        if self._project_combo.count():
            self._on_project_changed(self._project_combo.currentData() or '')

    def set_stages(self, stage_names: list[str]) -> None:
        """Vul de fase-combobox.

        Parameters
        ----------
        stage_names: Lijst van fase-namen.
        """
        self._stage_combo.blockSignals(True)
        self._stage_combo.clear()
        for i, name in enumerate(stage_names):
            self._stage_combo.addItem(name)
        self._stage_combo.blockSignals(False)

    def active_project_key(self) -> str | None:
        """Geef de base_name van het geselecteerde project."""
        return self._project_combo.currentData()

    def active_stage_index(self) -> int:
        """Geef de index van de geselecteerde fase."""
        return max(0, self._stage_combo.currentIndex())

''',

    'ui/file_list_widget.py': r'''
"""Widget voor de lijst van geÃ¯mporteerde bestanden."""

from __future__ import annotations
from PyQt6.QtWidgets import QListWidget, QListWidgetItem, QAbstractItemView
from PyQt6.QtCore import Qt


class FileListWidget(QListWidget):
    """QListWidget die geÃ¯mporteerde D-Sheet bestanden toont."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setSelectionMode(QAbstractItemView.SelectionMode.NoSelection)
        self.setAlternatingRowColors(True)
        self.setMaximumHeight(200)

    def set_files(self, files: dict[str, str]) -> None:
        """Vul de lijst met bestandsnamen.

        Parameters
        ----------
        files: Dict filename â†’ raw text.
        """
        self.clear()
        for name in sorted(files.keys()):
            item = QListWidgetItem(name)
            item.setToolTip(f'{len(files[name]):,} tekens')
            self.addItem(item)

''',

    'ui/controls_panel.py': r'''
"""Viewport- en renderschaalbediening."""

from __future__ import annotations
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QGroupBox, QLabel,
    QCheckBox, QDoubleSpinBox, QPushButton
)
from PyQt6.QtCore import pyqtSignal

from app.settings import RenderSettings, ViewportSettings


class ControlsPanel(QWidget):
    """Panel met viewport- en renderinstellingen.

    Signals
    -------
    viewport_changed:  Viewport-instellingen zijn gewijzigd.
    render_changed:    Render-instellingen zijn gewijzigd.
    zoom_in:           Zoom-in knop ingedrukt.
    zoom_out:          Zoom-out knop ingedrukt.
    zoom_reset:        Auto-reset knop ingedrukt.
    """

    viewport_changed = pyqtSignal(object)   # ViewportSettings
    render_changed = pyqtSignal(object)     # RenderSettings
    zoom_in = pyqtSignal()
    zoom_out = pyqtSignal()
    zoom_reset = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(4, 4, 4, 4)

        # --- Viewport ---
        vp_box = QGroupBox('Viewport')
        vp_layout = QVBoxLayout(vp_box)

        self._auto_check = QCheckBox('Automatisch bereik')
        self._auto_check.setChecked(True)
        self._auto_check.toggled.connect(self._on_viewport_change)
        vp_layout.addWidget(self._auto_check)

        def _spin(lo=-9999.0, hi=9999.0, val=0.0, step=0.5) -> QDoubleSpinBox:
            s = QDoubleSpinBox()
            s.setRange(lo, hi)
            s.setValue(val)
            s.setSingleStep(step)
            s.setDecimals(1)
            return s

        for label_text, attr, default in [
            ('X min', '_xmin', -10.0), ('X max', '_xmax', 10.0),
            ('Y min', '_ymin', -10.0), ('Y max', '_ymax', 5.0),
        ]:
            row = QWidget()
            row_l = QHBoxLayout(row)
            row_l.setContentsMargins(0, 0, 0, 0)
            row_l.addWidget(QLabel(label_text))
            spin = _spin(val=default)
            spin.valueChanged.connect(self._on_viewport_change)
            setattr(self, attr, spin)
            row_l.addWidget(spin)
            vp_layout.addWidget(row)

        zoom_row = QWidget()
        zoom_l = QHBoxLayout(zoom_row)
        zoom_l.setContentsMargins(0, 0, 0, 0)
        for label, signal in [('Zoom in', self.zoom_in),
                               ('Zoom uit', self.zoom_out),
                               ('Auto-reset', self.zoom_reset)]:
            btn = QPushButton(label)
            btn.clicked.connect(signal.emit)
            zoom_l.addWidget(btn)
        vp_layout.addWidget(zoom_row)
        layout.addWidget(vp_box)

        # --- Render instellingen ---
        rs_box = QGroupBox('Renderinstellingen')
        rs_layout = QVBoxLayout(rs_box)

        for label_text, attr, default, tooltip in [
            ('Uniform 10kPa = m', '_uni_scale', 0.5,
             'Hoogte belastingblok: m per 10 kPa'),
            ('Normaal 10kN/m = m', '_norm_scale', 0.5,
             'Breedte normaalkrachtdiagram: m per 10 kN/m'),
            ('H-last schaal (m)', '_hload_scale', 2.0, 'Stemlengthe horizontale lijnlast'),
            ('Moment radius (m)', '_mom_radius', 1.0, 'Straal momentensymbool'),
            ('Waterpeil schaal', '_waterpeil_schaal', 1.0,
             'Schalfactor golflijntjes en verticale stap waterpeils-symbool'),
            ('Maaiveld schaal', '_maaiveld_schaal', 1.0,
             'Schalfactor driehoekbreedte maaiveld-symbool'),
        ]:
            row = QWidget()
            row_l = QHBoxLayout(row)
            row_l.setContentsMargins(0, 0, 0, 0)
            lbl = QLabel(label_text)
            lbl.setToolTip(tooltip)
            row_l.addWidget(lbl)
            spin = _spin(lo=0.01, hi=100.0, val=default, step=0.1)
            spin.setToolTip(tooltip)
            spin.valueChanged.connect(self._on_render_change)
            setattr(self, attr, spin)
            row_l.addWidget(spin)
            rs_layout.addWidget(row)

        layout.addWidget(rs_box)
        layout.addStretch()

    def _on_viewport_change(self) -> None:
        self.viewport_changed.emit(self.get_viewport_settings())

    def _on_render_change(self) -> None:
        self.render_changed.emit(self.get_render_settings())

    def get_viewport_settings(self) -> ViewportSettings:
        """Lees de huidige viewport-instellingen uit de UI."""
        return ViewportSettings(
            auto=self._auto_check.isChecked(),
            x_min=self._xmin.value(),
            x_max=self._xmax.value(),
            y_min=self._ymin.value(),
            y_max=self._ymax.value(),
        )

    def get_render_settings(self) -> RenderSettings:
        """Lees de huidige render-instellingen uit de UI."""
        return RenderSettings(
            uniform_meters_per_10kpa=self._uni_scale.value(),
            normal_meters_per_10knm=self._norm_scale.value(),
            hload_scale=self._hload_scale.value(),
            moment_radius_meters=self._mom_radius.value(),
            waterpeil_schaal=self._waterpeil_schaal.value(),
            maaiveld_schaal=self._maaiveld_schaal.value(),
        )

    def set_viewport(self, vp: ViewportSettings) -> None:
        """Vul de UI in vanuit een ViewportSettings object.

        Parameters
        ----------
        vp: Te tonen viewport-instellingen.
        """
        self._auto_check.blockSignals(True)
        self._auto_check.setChecked(vp.auto)
        self._auto_check.blockSignals(False)
        for spin, val in [(self._xmin, vp.x_min), (self._xmax, vp.x_max),
                           (self._ymin, vp.y_min), (self._ymax, vp.y_max)]:
            spin.blockSignals(True)
            spin.setValue(val)
            spin.blockSignals(False)

''',

    'ui/layer_table.py': r'''
"""QTableWidget voor het tonen van grondlagen per profiel."""

from __future__ import annotations
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QLabel, QTableWidget, QTableWidgetItem, QHeaderView
)
from PyQt6.QtGui import QColor
from PyQt6.QtCore import Qt

from parsers.models import Project, Stage, SoilProfile
from ui.table_styles import REPORT_QTABLE_STYLE
from utils.color_utils import rgb_string_to_tuple


class LayerTableWidget(QWidget):
    """Toont grondlagen van links- en rechtsprofiel in een tabel."""

    def __init__(self, parent=None):
        super().__init__(parent)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        self._label = QLabel('Geen profiel')
        layout.addWidget(self._label)

        self._table = QTableWidget(0, 5)
        self._table.setHorizontalHeaderLabels(['Nr', 'Materiaal', 'Bov. [m]', 'Ond. [m]', 'Kleur'])
        self._table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self._table.setAlternatingRowColors(True)
        self._table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._table.verticalHeader().setVisible(False)
        self._table.setShowGrid(True)
        self._table.setStyleSheet(REPORT_QTABLE_STYLE)
        layout.addWidget(self._table)

    def update_table(
        self,
        project: Project | None,
        stage: Stage | None,
        left_profile: SoilProfile | None,
        right_profile: SoilProfile | None,
    ) -> None:
        """Vul de tabel met laaginformatie.

        Parameters
        ----------
        project:       Huidig project.
        stage:         Actieve bouwfase.
        left_profile:  Linkerprofiel (kan None zijn).
        right_profile: Rechterprofiel (kan None zijn).
        """
        self._table.setRowCount(0)
        if not project or (not left_profile and not right_profile):
            self._label.setText('Geen profiel gevonden.')
            return

        rows: list[tuple] = []
        for side_label, profile in [('L', left_profile), ('R', right_profile)]:
            if not profile:
                continue
            layers = profile.layers
            for i, layer in enumerate(layers):
                bottom = str(layers[i + 1].level) if i + 1 < len(layers) else '...'
                color_str = project.soil_color_map.get(layer.material, 'rgb(220,220,220)')
                rows.append((side_label, layer.nr, layer.material,
                              str(layer.level), bottom, color_str))

        self._table.setRowCount(len(rows))
        self._label.setText(
            f'Links: {left_profile.name if left_profile else "-"}  |  '
            f'Rechts: {right_profile.name if right_profile else "-"}'
        )
        for row_idx, (side, nr, material, top_v, bottom_v, color_str) in enumerate(rows):
            self._table.setItem(row_idx, 0, QTableWidgetItem(f'{side}{nr}'))
            self._table.setItem(row_idx, 1, QTableWidgetItem(material))
            self._table.setItem(row_idx, 2, QTableWidgetItem(top_v))
            self._table.setItem(row_idx, 3, QTableWidgetItem(bottom_v))
            color_item = QTableWidgetItem('')
            try:
                r, g, b = rgb_string_to_tuple(color_str)
                color_item.setBackground(QColor(int(r * 255), int(g * 255), int(b * 255)))
            except Exception:
                pass
            self._table.setItem(row_idx, 4, color_item)

''',

    'ui/preview_window.py': r'''
"""Zwevend Word-preview venster voor D-Sheet Dashboard."""

from __future__ import annotations
from datetime import datetime

from PyQt6.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QTextBrowser,
)
from PyQt6.QtCore import QSettings


class WordPreviewWindow(QMainWindow):
    """Zwevend venster dat een HTML-rapportweergave toont in QTextBrowser.

    Het venster is bewust 'dom': het ontvangt alleen een HTML-string via
    set_html() en heeft geen directe toegang tot AppState of controllers.
    Positie en grootte worden onthouden via QSettings.
    """

    def __init__(self, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        self.setWindowTitle('Word Preview')
        self.resize(720, 900)
        self._build()
        self._herstel_geometrie()

    def _build(self) -> None:
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # â”€â”€ Statusbalk â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        status_balk = QWidget()
        status_balk.setStyleSheet(
            'background: #f0f4f7; border-bottom: 1px solid #c4d4e0;'
        )
        status_layout = QHBoxLayout(status_balk)
        status_layout.setContentsMargins(10, 4, 10, 4)
        status_layout.setSpacing(0)

        self._count_label = QLabel('Geen secties')
        self._count_label.setStyleSheet(
            'font-size: 10px; color: #5a7a8a; '
            'font-family: "Segoe UI", sans-serif;'
        )
        self._tijd_label = QLabel('')
        self._tijd_label.setStyleSheet(
            'font-size: 10px; color: #999; font-style: italic; '
            'font-family: "Segoe UI", sans-serif;'
        )

        status_layout.addWidget(self._count_label)
        status_layout.addStretch()
        status_layout.addWidget(self._tijd_label)
        layout.addWidget(status_balk)

        # â”€â”€ Preview-browser â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
        self._browser = QTextBrowser()
        self._browser.setOpenLinks(False)
        self._browser.setStyleSheet('border: none; background: white;')
        layout.addWidget(self._browser, stretch=1)

    # ------------------------------------------------------------------
    # Publieke API
    # ------------------------------------------------------------------

    def set_html(self, html: str, sectie_count: int = 0) -> None:
        """Toon nieuwe HTML-inhoud en werk de statusbalk bij.

        Parameters
        ----------
        html:
            Volledige HTML-string voor QTextBrowser.setHtml().
        sectie_count:
            Aantal geselecteerde secties voor de statusregel.
        """
        self._browser.setHtml(html)
        enkelvoud = sectie_count == 1
        self._count_label.setText(
            f'{sectie_count} sectie geselecteerd'
            if enkelvoud else
            f'{sectie_count} secties geselecteerd'
        )
        self._tijd_label.setText(
            f'â†» Bijgewerkt: {datetime.now().strftime("%H:%M")}'
        )

    # ------------------------------------------------------------------
    # Geometrie-persistentie
    # ------------------------------------------------------------------

    def _herstel_geometrie(self) -> None:
        instellingen = QSettings('DKIB', 'DSheetDashboard')
        geom = instellingen.value('preview_window/geometry')
        if geom:
            self.restoreGeometry(geom)

    def closeEvent(self, event) -> None:  # type: ignore[override]
        instellingen = QSettings('DKIB', 'DSheetDashboard')
        instellingen.setValue('preview_window/geometry', self.saveGeometry())
        super().closeEvent(event)

''',

    'exporters/excel_exporter.py': r'''
"""ExcelExporter â€” exporteert een ReportPackage naar Excel (.xlsx).

JSON-sidecar formaat (naast .xltx sjabloon, zelfde naam + .map.json):
{
  "metadata": {
    "project_name": {"sheet": "Voorblad", "cell": "B3"},
    "title":        {"sheet": "Voorblad", "cell": "B5"}
  },
  "sections": {
    "sheet_piling": "Damwand",
    "moment_max":   "Resultaten"
  }
}
"""

from __future__ import annotations
import json
from pathlib import Path

import openpyxl

from reporting.models import ReportPackage, ReportSection


class ExcelExporter:
    """Schrijft een ReportPackage naar een .xlsx-bestand met openpyxl."""

    def export(self, package: ReportPackage, template_path: str | None,
               output_path: str) -> str | None:
        """Exporteer naar Excel.

        Returns:
            None bij succes, foutmelding (str) bij een uitzondering.
        """
        try:
            mapping = self._load_mapping(template_path)

            if template_path and Path(template_path).exists():
                wb = openpyxl.load_workbook(template_path)
            else:
                wb = openpyxl.Workbook()
                if wb.active:
                    wb.active.title = 'Metadata'

            if mapping:
                self._write_with_mapping(wb, package, mapping)
            else:
                self._write_metadata(wb, package)
                selected_ids = {
                    i.source_ref for i in package.selected_items
                    if i.source_ref and i.included_excel
                }
                # extra_sections worden alleen gefilterd als er een overeenkomend
                # ReportItem met source_ref in selected_items bestaat.
                all_sections = (
                    package.input_sections
                    + package.result_sections
                    + package.extra_sections
                )
                for sec in all_sections:
                    if selected_ids and sec.id not in selected_ids:
                        continue
                    self._write_section(wb, sec)

            wb.template = False  # voorkomt opslaan als .xltx bij template-input
            wb.save(output_path)
            return None
        except Exception as exc:
            return str(exc)

    # ------------------------------------------------------------------
    # JSON-sidecar laden
    # ------------------------------------------------------------------

    def _load_mapping(self, template_path: str | None) -> dict | None:
        if not template_path:
            return None
        sidecar = Path(template_path).with_suffix('').with_suffix('.map.json')
        if not sidecar.exists():
            # probeer ook .xlsx.map.json
            sidecar2 = Path(str(template_path) + '.map.json')
            if not sidecar2.exists():
                return None
            sidecar = sidecar2
        try:
            return json.loads(sidecar.read_text(encoding='utf-8'))
        except Exception:
            return None

    # ------------------------------------------------------------------
    # Schrijven met sidecar-mapping
    # ------------------------------------------------------------------

    def _write_with_mapping(self, wb, package: ReportPackage, mapping: dict) -> None:
        """Vul template-cellen in via sidecar-mapping."""
        # Metadata-velden
        meta_map = mapping.get('metadata', {})
        md = package.metadata
        for attr, loc in meta_map.items():
            value = getattr(md, attr, '') or ''
            sheet_name = loc.get('sheet', '')
            cell = loc.get('cell', '')
            if sheet_name in wb.sheetnames and cell:
                wb[sheet_name][cell] = value

        # Secties: schrijf naar benoemde sheets als ze in mapping staan
        sec_map = mapping.get('sections', {})
        all_sections = (
            package.input_sections
            + package.result_sections
            + package.extra_sections
        )
        for sec in all_sections:
            sheet_name = sec_map.get(sec.id)
            if sheet_name and sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                self._append_section_to_sheet(ws, sec)
            else:
                # Geen mapping â†’ schrijf als nieuw werkblad
                self._write_section(wb, sec)

    def _append_section_to_sheet(self, ws, section: ReportSection) -> None:
        """Voeg sectie-data toe aan een bestaand werkblad (achter bestaande data)."""
        start_row = ws.max_row + 2
        row = start_row
        for f in section.fields:
            ws.cell(row=row, column=1, value=f.label)
            val = f'{f.value} {f.unit}'.strip() if f.unit else f.value
            ws.cell(row=row, column=2, value=val)
            row += 1
        for table in section.tables:
            row += 1
            ws.cell(row=row, column=1, value=table.title)
            row += 1
            for col, header in enumerate(table.columns, start=1):
                ws.cell(row=row, column=col, value=header)
            row += 1
            for data_row in table.rows:
                for col, cell in enumerate(data_row, start=1):
                    ws.cell(row=row, column=col, value=cell)
                row += 1
        for tb in section.text_blocks:
            row += 1
            ws.cell(row=row, column=1, value=tb.effective_text)
            row += 1

    # ------------------------------------------------------------------
    # Standaard schrijven (geen mapping)
    # ------------------------------------------------------------------

    def _write_metadata(self, wb, package: ReportPackage) -> None:
        ws = (wb['Metadata'] if 'Metadata' in wb.sheetnames
              else wb.create_sheet('Metadata'))
        md = package.metadata
        rows = [
            ('Projectnaam',    md.project_name),
            ('Ordernummer',    md.order_number),
            ('Locatie',        md.location),
            ('Fase',           md.phase),
            ('Opdrachtgever',  md.client),
            ('Titel',          md.title),
            ('Revisie',        md.revision),
            ('Auteur',         md.author),
            ('Datum',          md.date),
        ]
        for r, (label, value) in enumerate(rows, start=1):
            ws.cell(row=r, column=1, value=label)
            ws.cell(row=r, column=2, value=value)

    def _write_section(self, wb, section: ReportSection) -> None:
        sheet_name = section.title[:31]
        if sheet_name in wb.sheetnames:
            sheet_name = sheet_name[:28] + '...'
        ws = wb.create_sheet(sheet_name)

        row = 1
        for f in section.fields:
            ws.cell(row=row, column=1, value=f.label)
            val = f'{f.value} {f.unit}'.strip() if f.unit else f.value
            ws.cell(row=row, column=2, value=val)
            row += 1

        for table in section.tables:
            if row > 1:
                row += 1
            ws.cell(row=row, column=1, value=table.title)
            row += 1
            for col, header in enumerate(table.columns, start=1):
                ws.cell(row=row, column=col, value=header)
            row += 1
            for data_row in table.rows:
                for col, cell in enumerate(data_row, start=1):
                    ws.cell(row=row, column=col, value=cell)
                row += 1

        for tb in section.text_blocks:
            if row > 1:
                row += 1
            ws.cell(row=row, column=1, value='Beschrijving')
            ws.cell(row=row, column=2, value=tb.effective_text)
            row += 1

''',

    'exporters/word_exporter.py': r'''
"""WordExporter â€” exporteert een ReportPackage naar Word (.docx).

JSON-sidecar formaat (naast .dotx sjabloon, zelfde naam + .map.json):
{
  "metadata": {
    "project_name": "bookmark_project",
    "title":        "bookmark_title"
  },
  "sections": {
    "sheet_piling": "Sectie 2.1 Damwand",
    "moment_max":   "Sectie 3.1 Momenten"
  }
}
Sleutels in 'metadata' zijn bladwijzernamen in het template.
Sleutels in 'sections' zijn koppen waaronder de secties worden ingevoegd.
"""

from __future__ import annotations
import io
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

from reporting.models import FaseInvoerSectie, ReportPackage, ReportSection
from reporting.figure_renderer import render_figuur

_DOTX_CONTENT_TYPE = (
    'application/vnd.openxmlformats-officedocument'
    '.wordprocessingml.template.main+xml'
)
_DOCX_CONTENT_TYPE = (
    'application/vnd.openxmlformats-officedocument'
    '.wordprocessingml.document.main+xml'
)


class WordExporter:
    """Schrijft een ReportPackage naar een .docx-bestand met python-docx."""

    def export(self, package: ReportPackage, template_path: str | None,
               output_path: str, project=None) -> str | None:
        """Exporteer naar Word.

        Returns:
            None bij succes, foutmelding (str) bij een uitzondering.
        """
        try:
            mapping = self._load_mapping(template_path)

            if template_path and Path(template_path).exists():
                doc = self._open_template(template_path)
            else:
                doc = Document()
            self._apply_theme_typography(doc)

            if mapping:
                self._write_with_mapping(doc, package, mapping, project)
            else:
                self._write_metadata(doc, package)
                selected_ids = {
                    i.source_ref for i in package.selected_items
                    if i.source_ref and i.included_word
                }
                # extra_sections worden alleen gefilterd als er een overeenkomend
                # ReportItem met source_ref in selected_items bestaat.
                all_sections = (
                    package.input_sections
                    + package.result_sections
                    + package.extra_sections
                )
                for sec in all_sections:
                    if selected_ids and sec.id not in selected_ids:
                        continue
                    self._write_section(doc, sec, project)

            doc.save(output_path)
            return None
        except Exception as exc:
            return str(exc)

    # ------------------------------------------------------------------
    # Sjabloon openen (.dotx Ã©n .docx)
    # ------------------------------------------------------------------

    def _open_template(self, path: str) -> Document:
        """Open een Word-sjabloon als bewerkbaar Document.

        Een .dotx-bestand heeft een ander content type dan .docx, waardoor
        python-docx het weigert. We passen het content type in-memory aan
        zodat Document() het accepteert â€” zonder tijdelijke bestanden.

        Parameters
        ----------
        path: Pad naar het sjabloonbestand (.dotx of .docx).
        """
        if Path(path).suffix.lower() != '.dotx':
            return Document(path)

        with open(path, 'rb') as f:
            data = f.read()

        invoer = io.BytesIO(data)
        uitvoer = io.BytesIO()
        with zipfile.ZipFile(invoer, 'r') as zin, \
                zipfile.ZipFile(uitvoer, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                inhoud = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    inhoud = inhoud.replace(
                        _DOTX_CONTENT_TYPE.encode(),
                        _DOCX_CONTENT_TYPE.encode(),
                    )
                zout.writestr(item, inhoud)

        uitvoer.seek(0)
        return Document(uitvoer)

    def _apply_theme_typography(self, doc) -> None:
        """Pas thematekstgrootte toe op standaardtekst buiten tabellen."""
        from exporters.word_hoofdstuk_exporter import _eerste_fontfamilie
        from ui import table_styles

        stijl = doc.styles['Normal']
        stijl.font.name = _eerste_fontfamilie(table_styles.TABLE_FONT)
        stijl.font.size = Pt(table_styles.WORD_BODY_TEXT_SIZE)

    # ------------------------------------------------------------------
    # JSON-sidecar laden
    # ------------------------------------------------------------------

    def _load_mapping(self, template_path: str | None) -> dict | None:
        if not template_path:
            return None
        sidecar = Path(template_path).with_suffix('').with_suffix('.map.json')
        if not sidecar.exists():
            sidecar2 = Path(str(template_path) + '.map.json')
            if not sidecar2.exists():
                return None
            sidecar = sidecar2
        try:
            return json.loads(sidecar.read_text(encoding='utf-8'))
        except Exception:
            return None

    # ------------------------------------------------------------------
    # Schrijven met sidecar-mapping
    # ------------------------------------------------------------------

    def _write_with_mapping(
        self, doc, package: ReportPackage, mapping: dict, project=None
    ) -> None:
        """Vul template-bladwijzers in en voeg secties toe na genoemde koppen."""
        # Metadata via bladwijzers
        meta_map = mapping.get('metadata', {})
        md = package.metadata
        for attr, bookmark_name in meta_map.items():
            value = getattr(md, attr, '') or ''
            self._fill_bookmark(doc, bookmark_name, value)

        # Secties
        sec_map = mapping.get('sections', {})
        all_sections = (
            package.input_sections
            + package.result_sections
            + package.extra_sections
        )
        for sec in all_sections:
            heading_text = sec_map.get(sec.id)
            if heading_text:
                self._insert_after_heading(doc, heading_text, sec, project)
            else:
                self._write_section(doc, sec, project)

    def _fill_bookmark(self, doc, bookmark_name: str, value: str) -> None:
        """Vervang tekst van een bladwijzer in het document."""
        from docx.oxml.ns import qn
        for para in doc.paragraphs:
            for bm in para._element.findall(f'.//{qn("w:bookmarkStart")}'):
                if bm.get(qn('w:name')) == bookmark_name:
                    # Verwijder bestaande runs in de paragraaf en voeg waarde toe
                    for run in para.runs:
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = value
                    else:
                        para.add_run(value)
                    return

    def _insert_after_heading(self, doc, heading_text: str,
                               section: ReportSection, project=None) -> None:
        """Voeg sectie-content toe na de paragraaf met de gegeven koptekst."""
        from docx.oxml import OxmlElement
        target = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == heading_text.strip():
                target = para._element
                break
        if target is None:
            # Kop niet gevonden â†’ aan het einde toevoegen
            self._write_section(doc, section)
            return
        # Bouw sectie-XML op en voeg in na de kop
        temp_doc = doc.__class__()
        self._write_section(temp_doc, section, project)
        parent = target.getparent()
        idx = list(parent).index(target) + 1
        for elem in list(temp_doc.element.body)[:-1]:  # skip sectPr
            parent.insert(idx, elem)
            idx += 1

    # ------------------------------------------------------------------
    # Standaard schrijven (geen mapping)
    # ------------------------------------------------------------------

    def _write_metadata(self, doc, package: ReportPackage) -> None:
        md = package.metadata
        doc.add_heading('Rapportgegevens', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = 'Veld'
        table.rows[0].cells[1].text = 'Waarde'
        rows = [
            ('Projectnaam',   md.project_name),
            ('Ordernummer',   md.order_number),
            ('Locatie',       md.location),
            ('Fase',          md.phase),
            ('Opdrachtgever', md.client),
            ('Titel',         md.title),
            ('Revisie',       md.revision),
            ('Auteur',        md.author),
            ('Datum',         md.date),
        ]
        for label, value in rows:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value or ''

    def _write_section(self, doc, section: ReportSection, project=None) -> None:
        if isinstance(section, FaseInvoerSectie):
            from exporters.word_hoofdstuk_exporter import WordHoofdstukExporter

            WordHoofdstukExporter()._schrijf_fase_sectie(doc, section, project)
            return
        if section.id == 'damwand_gegevens':
            from exporters.word_hoofdstuk_exporter import WordHoofdstukExporter

            WordHoofdstukExporter()._schrijf_damwandgegevens_sectie(doc, section)
            return

        doc.add_heading(section.title, level=2)

        for f in section.fields:
            val = f'{f.value} {f.unit}'.strip() if f.unit else f.value
            doc.add_paragraph(f'{f.label}: {val}')

        for table_data in section.tables:
            doc.add_paragraph(table_data.title, style='Intense Quote')
            if not table_data.columns:
                continue
            tbl = doc.add_table(rows=1, cols=len(table_data.columns))
            for col, header in enumerate(table_data.columns):
                tbl.rows[0].cells[col].text = header
            for data_row in table_data.rows:
                row = tbl.add_row()
                for col, cell in enumerate(data_row):
                    if col < len(row.cells):
                        row.cells[col].text = cell

        for tb in section.text_blocks:
            doc.add_paragraph(tb.effective_text)

        for groep in section.image_groups:
            self._write_image_group(doc, groep, project)

        for img_req in section.images:
            self._write_image(doc, img_req, project)

    def _write_image(self, doc, img_req, project) -> None:
        """Render een figuurverzoek en voeg het toe aan het Word-document."""
        if project is None:
            doc.add_paragraph(f'[Figuur: {img_req.caption}]')
            return
        png_bytes = render_figuur(img_req, project)
        if png_bytes:
            doc.add_picture(io.BytesIO(png_bytes), width=Cm(16))
            if img_req.caption:
                doc.add_paragraph(img_req.caption, style='Caption')
        else:
            doc.add_paragraph(f'[Figuur niet beschikbaar: {img_req.caption}]')

    def _write_image_group(self, doc, groep, project) -> None:
        """Schrijf een figuurgroep als Word-tabel met kop, figuur en bron."""
        if not groep.headers:
            return
        if groep.title:
            doc.add_paragraph(groep.title, style='Intense Quote')

        tbl = doc.add_table(rows=3, cols=len(groep.headers))
        for col, header in enumerate(groep.headers):
            tbl.rows[0].cells[col].text = header

        for col, img_req in enumerate(groep.images):
            cell = tbl.rows[1].cells[col]
            if img_req is None:
                cell.text = '-'
                continue
            if project is None:
                cell.text = f'[Figuur: {img_req.caption or img_req.figure_key}]'
                continue
            png_bytes = render_figuur(img_req, project)
            if not png_bytes:
                cell.text = '[Figuur niet beschikbaar]'
                continue
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(io.BytesIO(png_bytes), width=Cm(5.2))

        for col, footer in enumerate(groep.footers):
            if col < len(tbl.rows[2].cells):
                tbl.rows[2].cells[col].text = footer

''',

    'reporting/builders/html_preview_builder.py': r'''
"""HtmlPreviewBuilder â€” genereert een HTML-string uit een ReportPackage."""

from __future__ import annotations

import base64

from reporting.models import ReportPackage, ReportSection, ReportField, ReportTable
from reporting.figure_renderer import render_figuur

# â”€â”€ Kleurconstanten (consistent met app-stijl) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
_HDR_BG   = '#147ACF'
_HDR_FG   = '#ffffff'
_SUB_BG   = '#147ACF'
_SUB_FG   = '#ffffff'
_ODD_BG   = '#ffffff'
_EVEN_BG  = '#f2f2f2'
_SEP      = '#000000'
_LABEL    = '#000000'
_VALUE    = '#000000'
_FONT     = '"Segoe UI", "Helvetica Neue", Arial, sans-serif'

_CSS = f"""
  body {{ font-family: {_FONT}; font-size: 12px; color: {_VALUE};
          margin: 0; padding: 16px; background: #ffffff; }}
  h1   {{ font-size: 15px; font-weight: 700; color: {_HDR_BG};
          border-bottom: 2px solid {_HDR_BG}; padding-bottom: 6px;
          margin-bottom: 16px; }}
  h2   {{ font-size: 12px; font-weight: 700; color: {_SUB_BG};
          margin: 18px 0 6px 0; padding: 5px 10px;
          background: #eaf2f8; border-left: 3px solid {_SUB_BG}; }}
  table {{ width: 100%; border-collapse: collapse; margin-bottom: 12px;
           border: 1px solid {_SEP}; }}
  th    {{ background: {_HDR_BG}; color: {_HDR_FG}; font-size: 11px;
           font-weight: 600; padding: 5px 10px; text-align: left;
           border: 1px solid {_SEP}; }}
  td    {{ padding: 5px 10px; border: 1px solid {_SEP};
           font-size: 11px; }}
  tr.odd  td {{ background: {_ODD_BG}; }}
  tr.even td {{ background: {_EVEN_BG}; }}
  td.label {{ color: {_LABEL}; font-weight: 500; width: 45%; }}
  td.value {{ text-align: right; }}
  td.unit  {{ color: {_VALUE}; font-size: 10px; width: 20%; }}
  p.tekst  {{ font-size: 11px; color: #3d4f5c; margin: 4px 0 10px 0;
              line-height: 1.6; }}
  p.caption {{ font-size: 10px; color: #666666; margin: 2px 0 10px 0; }}
  p.leeg   {{ color: #a0b4c2; font-style: italic; padding: 20px 0; }}
  img.figuur {{ max-width: 100%; margin: 8px 0 2px 0; }}
  img.figuur-cel {{ width: 100%; max-width: 100%; margin: 4px 0; }}
  table.figuurgroep td {{ vertical-align: top; text-align: center; }}
  table.figuurgroep .bron {{ font-size: 10px; color: #555555; }}
  .inline-wrap td {{ vertical-align: top; padding-right: 16px; }}
  .inline-wrap {{ border-collapse: separate; border-spacing: 0; width: auto;
                  margin-bottom: 12px; border: none; }}
  .inline-wrap table {{ width: auto; min-width: 160px; margin-bottom: 0; }}
"""


class HtmlPreviewBuilder:
    """Zet een ReportPackage om naar een HTML-string voor QTextBrowser."""

    def build(self, package: ReportPackage, project=None) -> str:
        """Genereer HTML-string voor de geselecteerde secties.

        Parameters
        ----------
        package:
            Rapportpakket met invoer-, resultaat- en extra-secties en de selectielijst.
        project:
            Actief project voor het renderen van figuren; ``None`` laat figuren weg.

        Returns
        -------
        str
            Volledige HTML-string geschikt voor QTextBrowser.setHtml().
        """
        titel = package.metadata.project_name or 'Rapport'

        alle_secties: dict[str, ReportSection] = {
            s.id: s
            for s in (
                package.input_sections
                + package.result_sections
                + package.extra_sections
            )
        }

        secties: list[str] = []
        for item in package.selected_items:
            if not item.included_word:
                continue
            sec = alle_secties.get(item.source_ref)
            if sec is not None:
                secties.append(self._sectie_html(sec, project))

        body = (
            '\n'.join(secties)
            if secties
            else '<p class="leeg">Geen secties geselecteerd.</p>'
        )

        return (
            f'<!DOCTYPE html><html><head>'
            f'<meta charset="utf-8">'
            f'<style>{_CSS}</style>'
            f'</head><body>'
            f'<h1>{_esc(titel)}</h1>'
            f'{body}'
            f'</body></html>'
        )

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _sectie_html(self, sec: ReportSection, project=None) -> str:
        """Render Ã©Ã©n ReportSection als HTML-fragment."""
        delen: list[str] = [f'<h2>{_esc(sec.title)}</h2>']

        if sec.fields:
            delen.append(self._velden_html(sec.fields))

        # Inline tabellen naast elkaar in flex-container; overige tabellen gestapeld
        inline = [t for t in sec.tables if t.inline]
        normaal = [t for t in sec.tables if not t.inline]
        if inline:
            cellen = ''.join(
                f'<td>{self._tabel_html(t)}</td>' for t in inline
            )
            delen.append(
                f'<table class="inline-wrap"><tr>{cellen}</tr></table>'
            )
        for tabel in normaal:
            delen.append(self._tabel_html(tabel))

        for blok in sec.text_blocks:
            tekst = blok.effective_text
            if tekst:
                delen.append(f'<p class="tekst">{_esc(tekst)}</p>')

        for groep in sec.image_groups:
            groep_html = self._figuurgroep_html(groep, project)
            if groep_html:
                delen.append(groep_html)

        for img_req in sec.images:
            figuur_html = self._figuur_html(img_req, project)
            if figuur_html:
                delen.append(figuur_html)

        return '\n'.join(delen)

    def _figuur_html(self, img_req, project=None) -> str:
        """Render een figuur als base64 data-URI voor de HTML-preview."""
        if project is None:
            return ''
        png = render_figuur(img_req, project)
        if not png:
            return ''
        b64 = base64.b64encode(png).decode('ascii')
        caption = (
            f'<p class="caption">{_esc(img_req.caption)}</p>'
            if img_req.caption else ''
        )
        return f'<img class="figuur" src="data:image/png;base64,{b64}">{caption}'

    def _figuurgroep_html(self, groep, project=None) -> str:
        """Render een groep figuren als 3-rijen tabel."""
        if not groep.headers:
            return ''

        header = ''.join(f'<th>{_esc(kop)}</th>' for kop in groep.headers)
        figuur_cellen: list[str] = []
        for img_req in groep.images:
            if img_req is None or project is None:
                figuur_cellen.append('<td>-</td>')
                continue
            png = render_figuur(img_req, project)
            if not png:
                figuur_cellen.append('<td>-</td>')
                continue
            b64 = base64.b64encode(png).decode('ascii')
            figuur_cellen.append(
                f'<td><img class="figuur-cel" '
                f'src="data:image/png;base64,{b64}"></td>'
            )
        footer = ''.join(
            f'<td class="bron">{_esc(tekst)}</td>' for tekst in groep.footers
        )
        titel = (
            f'<p style="font-size:11px;font-weight:600;color:{_LABEL};'
            f'margin:8px 0 3px;">{_esc(groep.title)}</p>'
            if groep.title else ''
        )
        return (
            f'{titel}<table class="figuurgroep">'
            f'<tr>{header}</tr>'
            f'<tr>{"".join(figuur_cellen)}</tr>'
            f'<tr>{footer}</tr>'
            f'</table>'
        )

    def _velden_html(self, velden: list[ReportField]) -> str:
        """Render veld-rijen als HTML-tabel."""
        rijen = []
        for i, veld in enumerate(velden):
            klasse = 'odd' if i % 2 == 0 else 'even'
            unit_cel = f'<td class="unit">{_esc(veld.unit)}</td>' if veld.unit else '<td></td>'
            rijen.append(
                f'<tr class="{klasse}">'
                f'<td class="label">{_esc(veld.label)}</td>'
                f'<td class="value">{_esc(veld.value)}</td>'
                f'{unit_cel}'
                f'</tr>'
            )
        return f'<table>{"".join(rijen)}</table>'

    def _tabel_html(self, tabel: ReportTable) -> str:
        """Render een ReportTable als HTML-tabel met header."""
        seps = set(tabel.separator_before_cols)
        sep_style = f'border-left: 2px solid {_SEP};'

        def th(i: int, k: str) -> str:
            st = f' style="{sep_style}"' if i in seps else ''
            return f'<th{st}>{_esc(k)}</th>'

        def td(i: int, cel: str) -> str:
            st = f' style="{sep_style}"' if i in seps else ''
            return f'<td{st}>{_esc(cel)}</td>'

        if tabel.column_groups:
            # Eerste headerrij: groepkoppen met colspan (geen rowspan)
            col_idx = 0
            groep_cellen: list[str] = []
            for label, span in tabel.column_groups:
                st = ''
                if col_idx in seps:
                    st = f' style="text-align:center;{sep_style}"'
                elif label:
                    st = ' style="text-align:center;"'
                groep_cellen.append(
                    f'<th{st} colspan="{span}">{_esc(label)}</th>'
                )
                col_idx += span
            groep_rij = f'<tr>{"".join(groep_cellen)}</tr>'
            # Tweede headerrij: alle individuele kolomkoppen
            header = ''.join(th(i, k) for i, k in enumerate(tabel.columns))
            header_html = f'{groep_rij}<tr>{header}</tr>'
        else:
            header = ''.join(th(i, k) for i, k in enumerate(tabel.columns))
            header_html = f'<tr>{header}</tr>'

        rijen = []
        for row_i, rij in enumerate(tabel.rows):
            klasse = 'odd' if row_i % 2 == 0 else 'even'
            cellen = ''.join(td(i, cel) for i, cel in enumerate(rij))
            rijen.append(f'<tr class="{klasse}">{cellen}</tr>')
        if tabel.title:
            kop = f'<p style="font-size:11px;font-weight:600;color:{_LABEL};margin:8px 0 3px;">{_esc(tabel.title)}</p>'
        else:
            kop = ''
        return f'{kop}<table>{header_html}{"".join(rijen)}</table>'


def _esc(tekst: object) -> str:
    """Vervang HTML-speciale tekens door entiteiten.

    Parameters
    ----------
    tekst:
        Te escapen waarde; wordt eerst naar str geconverteerd.

    Returns
    -------
    str
        HTML-veilige string met speciale tekens vervangen door entiteiten.
    """
    return (
        str(tekst)
        .replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
        .replace('"', '&quot;')
    )

''',

    'app/report_controller.py': r'''
"""ReportController â€” orkestreert rapportagelaag voor D-Sheet Dashboard.

Heeft geen Qt-kennis. Retourneert primitieven of domeinobjecten.
"""

from __future__ import annotations

from app.state import AppState
from app.report_state import ReportState
from reporting.models import ReportItem, ReportSection, ReportPackage
from reporting.selection import ReportPlan
from reporting.builders.damwand_hoofdstuk_builder import DamwandHoofdstukBuilder
from reporting.builders.input_description_builder import InputDescriptionBuilder, DamwandCard
from reporting.builders.result_description_builder import ResultDescriptionBuilder
from reporting.builders.soil_table_builder import SoilTableBuilder
from exporters.excel_exporter import ExcelExporter
from exporters.word_hoofdstuk_exporter import WordHoofdstukExporter


class ReportController:
    """Applicatielaag voor rapportage: builders, plan, exporters, validatie."""

    def __init__(self, app_state: AppState, report_state: ReportState) -> None:
        self._app = app_state
        self._report = report_state
        self._damwand_builder = DamwandHoofdstukBuilder()
        self._input_builder = InputDescriptionBuilder()
        self._result_builder = ResultDescriptionBuilder()
        self._soil_builder = SoilTableBuilder()
        self._excel = ExcelExporter()

    # ------------------------------------------------------------------
    # Builders
    # ------------------------------------------------------------------

    def build_all_fase_cards(self):
        """Bouw FaseCard-lijst voor alle fases van het actieve project."""
        project = self._app.get_active_project()
        if not project:
            return []
        return self._input_builder.build_all_stages(project)

    def build_damwand_card(self) -> DamwandCard | None:
        """Bouw DamwandCard voor het actieve project."""
        project = self._app.get_active_project()
        if not project:
            return None
        return self._input_builder.build_damwand_card(project)

    def build_input_descriptions(self) -> list[ReportSection]:
        """Bouw invoerbeschrijvingssecties voor actief project/fase."""
        project = self._app.get_active_project()
        stage = self._app.get_active_stage()
        if not project or not stage:
            return []
        return self._input_builder.build(project, stage, self._report.overrides)

    def build_damwand_sections(self) -> list[ReportSection]:
        """Bouw damwand-invoersecties voor het actieve project.

        Returns
        -------
        list[ReportSection]
            Damwandgegevens als eerste sectie, gevolgd door fase-invoer.
        """
        project = self._app.get_active_project()
        if not project:
            return []
        return self._damwand_builder.build_input_sections(project)

    def build_result_descriptions(self) -> list[ReportSection]:
        """Bouw resultaatbeschrijvingssecties voor actief project/fase/stap."""
        project = self._app.get_active_project()
        if not project:
            return []
        return self._result_builder.build(
            project,
            self._app.active_output_stage_index,
            self._app.active_result_step,
            self._report.overrides,
        )

    def build_soil_sections(self) -> list[ReportSection]:
        """Bouw grondsoortentabelsecties voor het actieve project.

        Returns
        -------
        list[ReportSection]
            Ã‰Ã©n sectie per grondprofiel, lege lijst als er geen project is.
        """
        project = self._app.get_active_project()
        if not project:
            return []
        return self._soil_builder.build(project)

    # ------------------------------------------------------------------
    # Templates
    # ------------------------------------------------------------------

    def set_template_excel(self, path: str | None) -> None:
        """Sla het Excel-templatepad op."""
        self._report.template_excel = path or None

    def set_template_word(self, path: str | None) -> None:
        """Sla het Word-templatepad op."""
        self._report.template_word = path or None

    # ------------------------------------------------------------------
    # Overrides (stap 9)
    # ------------------------------------------------------------------

    def set_text_override(self, block_id: str, text: str) -> None:
        """Sla een handmatige tekstoverride op (leeg = verwijder override)."""
        if text.strip():
            self._report.overrides[block_id] = text
        else:
            self._report.overrides.pop(block_id, None)

    # ------------------------------------------------------------------
    # Plan
    # ------------------------------------------------------------------

    def get_plan(self) -> ReportPlan:
        """Geef het huidige rapportageplan terug."""
        return self._report.plan

    def auto_populate_plan(self) -> None:
        """Vul het rapportplan automatisch met secties van DamwandHoofdstukBuilder.

        Het bestaande plan wordt eerst leeggemaakt zodat de lijst altijd het
        actieve project weerspiegelt.
        """
        project = self._app.get_active_project()
        if not project:
            self._report.plan.items.clear()
            return
        self._report.plan.items.clear()
        secties = self._damwand_builder.build(project, None, None)
        gewenste_ids: list[str] = []
        for sec in secties:
            gewenste_ids.append(sec.id)
            self._report.plan.add_item(ReportItem(
                id=sec.id,
                kind=self._sectie_kind(sec.id),
                caption=sec.title,
                source_ref=sec.id,
            ))
        self._orden_plan_items(gewenste_ids)

    _RESULTAAT_IDS: frozenset[str] = frozenset({
        'anchor_forces', 'per_phase_summary', 'extremen_overzicht',
    })

    def _sectie_kind(self, sec_id: str) -> str:
        if sec_id in self._RESULTAAT_IDS:
            return 'resultaat'
        if sec_id.startswith('grondsoorten'):
            return 'grondsoorten'
        return 'invoer'

    def _orden_plan_items(self, gewenste_ids: list[str]) -> None:
        """Orden bestaande planitems volgens de opgegeven id-volgorde."""
        volgorde = {item_id: i for i, item_id in enumerate(gewenste_ids)}
        self._report.plan.items.sort(
            key=lambda item: (volgorde.get(item.id, len(volgorde)), item.order)
        )
        self._report.plan._renumber()

    # ------------------------------------------------------------------
    # Pakketbouw
    # ------------------------------------------------------------------

    def build_package(self) -> ReportPackage:
        """Bouw een ReportPackage op basis van huidige state."""
        input_secs = self.build_damwand_sections()
        result_secs = self.build_result_descriptions()
        soil_secs = self.build_soil_sections()
        pkg = self._report.plan.build_package(
            self._report.metadata, input_secs, result_secs,
            extra_sections=soil_secs,
        )
        pkg.template_excel = self._report.template_excel
        pkg.template_word = self._report.template_word
        return pkg

    def build_metadata(self) -> 'ReportMetadata':
        """Geef ReportMetadata op basis van de huidige rapport-state."""
        from reporting.models import ReportMetadata
        rs = self._report
        return ReportMetadata(
            project_name=getattr(rs, 'project_name', '') or '',
            client=getattr(rs, 'client', '') or '',
            author=getattr(rs, 'author', '') or '',
            date=getattr(rs, 'date', '') or '',
            revision=getattr(rs, 'revision', '') or '',
        )

    # ------------------------------------------------------------------
    # Export
    # ------------------------------------------------------------------

    def export_excel(self, output_path: str) -> str | None:
        """Exporteer naar Excel.

        Returns:
            None bij succes, foutmelding bij een fout.
        """
        package = self.build_package()
        return self._excel.export(package, self._report.template_excel, output_path)

    def export_word(self, output_path: str) -> str | None:
        """Exporteer naar Word via WordHoofdstukExporter.

        Returns
        -------
        str | None
            None bij succes, foutmelding bij een fout.
        """
        project = self._app.get_active_project()
        if not project:
            return 'Geen actief project geladen.'
        alle_secties = self._damwand_builder.build(project, None, None)
        geselecteerd = {
            item.source_ref for item in self._report.plan.items if item.included_word
        }
        secties = [s for s in alle_secties if s.id in geselecteerd] if geselecteerd else alle_secties
        metadata = self.build_metadata()
        template = (
            self._report.template_word
            or self._app.app_settings.word_template_path
            or 'templates/damwand_stijlen.docx'
        )
        return WordHoofdstukExporter().export(
            sections=secties,
            metadata=metadata,
            project=project,
            template_path=template,
            output_path=output_path,
        )

''',

    'parsers/__init__.py': r'''
"""Parser-pakket met plugin-registry voor D-Sheet bestandsformaten."""

from __future__ import annotations
from typing import Callable, Type

# Registry: extensie â†’ parser-callable
_registry: dict[str, Callable] = {}


def register_parser(extension: str, parser_callable: Callable) -> None:
    """Registreer een parser voor een bestandsextensie.

    Parameters
    ----------
    extension:        Bestandsextensie zonder punt (bijv. 'shi', 'plx').
    parser_callable:  Functie of klasse die (text: str, base_name: str) â†’ Project aanneemt.
    """
    _registry[extension.lower()] = parser_callable


def get_parser(extension: str) -> Callable | None:
    """Geef de geregistreerde parser terug voor de opgegeven extensie.

    Parameters
    ----------
    extension: Bestandsextensie zonder punt.

    Returns
    -------
    Callable | None  Parser-callable, of None als niet gevonden.
    """
    return _registry.get(extension.lower())


# Registreer de ingebouwde D-Sheet parser
from parsers.shi_parser import parse_project  # noqa: E402
from parsers.models import FileBundle         # noqa: E402

def _dsheet_parser(file_bundle: FileBundle, base_name: str):
    return parse_project(file_bundle, base_name)

register_parser('shi', _dsheet_parser)
register_parser('shd', _dsheet_parser)
register_parser('shs', _dsheet_parser)

''',

    'renderers/draw_helpers.py': r'''
"""Matplotlib-tekenhulpfuncties voor de D-Sheet renderer."""

from __future__ import annotations
import math
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.axes import Axes
from matplotlib.patches import Arc

from utils.color_utils import color_for_matplotlib


def draw_polygon_on_ax(
    ax: Axes,
    points: list[tuple[float, float]],
    face_color=None,
    edge_color=None,
    line_width: float = 1.0,
    alpha: float = 1.0,
    clip_on: bool = True,
) -> None:
    """Teken een gevuld polygoon op een matplotlib-assen.

    Parameters
    ----------
    ax:         Matplotlib Axes.
    points:     Lijst van (x, y) tuples in data-coÃ¶rdinaten.
    face_color: Vulkleur (None = transparant).
    edge_color: Randkleur (None = geen rand).
    line_width: Lijndikte.
    alpha:      Transparantie (0-1).
    clip_on:    Begrens tekenen tot het axes-gebied.
    """
    if not points or len(points) < 3:
        return
    xs, ys = zip(*points)
    poly = plt.Polygon(
        list(zip(xs, ys)),
        closed=True,
        facecolor=face_color if face_color is not None else 'none',
        edgecolor=edge_color if edge_color is not None else 'none',
        linewidth=line_width,
        alpha=alpha,
        clip_on=clip_on,
    )
    ax.add_patch(poly)


def fill_with_vertical_hatch(
    ax: Axes,
    points: list[tuple[float, float]],
    face_color=None,
    edge_color: str = '#333333',
    spacing: float = 0.3,
    clip_on: bool = True,
) -> None:
    """Vul een polygoon met verticale arceerlijnen (voor grondsoorten).

    Parameters
    ----------
    ax:         Matplotlib Axes.
    points:     Polygoon in data-coÃ¶rdinaten.
    face_color: Achtergrondkleur.
    edge_color: Kleur van de arceerlijnen.
    spacing:    Afstand tussen arceerlijnen in data-eenheden.
    clip_on:    Begrens tekenen tot het axes-gebied.
    """
    if not points or len(points) < 3:
        return
    draw_polygon_on_ax(ax, points, face_color=face_color, edge_color=edge_color,
                        line_width=0.8, clip_on=clip_on)
    xs, ys = zip(*points)
    x_min, x_max = min(xs), max(xs)
    x = x_min
    while x <= x_max + spacing:
        ax.plot([x, x], [min(ys) - 1, max(ys) + 1],
                color=edge_color, linewidth=0.7, clip_on=clip_on)
        x += spacing


def fill_with_surface_aligned_arrows(
    ax: Axes,
    top_pts: list[tuple[float, float]],
    bottom_pts: list[tuple[float, float]],
    face_color=None,
    edge_color: str = '#222222',
    spacing: float = 0.5,
    clip_on: bool = True,
) -> None:
    """Vul een belastingblok met neerwaartse pijlen langs het maaiveld.

    Parameters
    ----------
    ax:          Matplotlib Axes.
    top_pts:     Bovenpunten van het belastingblok (data-coÃ¶rdinaten).
    bottom_pts:  Onderpunten (maaiveldlijn).
    face_color:  Achtergrondkleur van het blok.
    edge_color:  Kleur van de pijlen.
    spacing:     Horizontale afstand tussen pijlen in data-eenheden.
    clip_on:     Begrens tekenen tot het axes-gebied.
    """
    if not top_pts or not bottom_pts:
        return
    all_pts = list(top_pts) + list(reversed(bottom_pts))
    draw_polygon_on_ax(ax, all_pts, face_color=face_color or 'white',
                        edge_color=edge_color, line_width=0.8, alpha=0.85, clip_on=clip_on)

    def interp_y(pts: list[tuple[float, float]], x: float) -> float:
        if len(pts) == 1:
            return pts[0][1]
        for i in range(len(pts) - 1):
            ax_x, ax_y = pts[i]
            bx, by = pts[i + 1]
            mn = min(ax_x, bx)
            mx = max(ax_x, bx)
            if mn - 1e-6 <= x <= mx + 1e-6:
                dx = bx - ax_x
                if abs(dx) < 1e-9:
                    return ax_y
                return ax_y + (x - ax_x) / dx * (by - ax_y)
        return pts[-1][1] if x > pts[-1][0] else pts[0][1]

    if len(bottom_pts) >= 2:
        x_start = bottom_pts[0][0]
        x_end = bottom_pts[-1][0]
        n_arrows = max(1, int(abs(x_end - x_start) / spacing))
        for k in range(n_arrows + 1):
            t = (k + 0.5) / (n_arrows + 1)
            bx = x_start + (x_end - x_start) * t
            by = interp_y(list(bottom_pts), bx)
            ty = interp_y(list(top_pts), bx)
            arrow_len = by - ty
            if abs(arrow_len) > 1e-6:
                ax.annotate(
                    '', xy=(bx, by), xytext=(bx, ty),
                    arrowprops=dict(arrowstyle='->', color=edge_color, lw=1.0),
                    clip_on=clip_on
                )


def fill_with_diagonal_hatch(
    ax: Axes,
    points: list[tuple[float, float]],
    face_color: str = 'white',
    edge_color: str = '#222222',
    spacing: float = 0.3,
    alpha: float = 0.7,
    clip_on: bool = True,
) -> None:
    """Vul een polygoon met diagonale arceerlijnen (voor normaalkrachten).

    Parameters
    ----------
    ax:         Matplotlib Axes.
    points:     Polygoon in data-coÃ¶rdinaten.
    face_color: Achtergrondkleur.
    edge_color: Kleur van de diagonale lijnen.
    spacing:    Afstand tussen lijnen.
    alpha:      Transparantie.
    clip_on:    Begrens tekenen tot het axes-gebied.
    """
    if not points or len(points) < 3:
        return
    draw_polygon_on_ax(ax, points, face_color=face_color, edge_color=edge_color,
                        line_width=0.8, alpha=alpha, clip_on=clip_on)
    xs, ys = zip(*points)
    x_min, x_max = min(xs), max(xs)
    y_min, y_max = min(ys), max(ys)
    span = (y_max - y_min) + 2
    x = x_min - span
    while x <= x_max + span:
        ax.plot([x, x + span], [y_max + 1, y_min - 1],
                color=edge_color, linewidth=0.7, clip_on=clip_on)
        x += spacing


def draw_moment_symbol(
    ax: Axes,
    x: float,
    y: float,
    clockwise: bool = True,
    color: str = '#111111',
    radius: float = 0.8,
) -> None:
    """Teken een gebogen pijl als momentensymbool.

    Parameters
    ----------
    ax:        Matplotlib Axes.
    x, y:      Middelpunt in data-coÃ¶rdinaten.
    clockwise: Richting van het moment.
    color:     Kleur van het symbool.
    radius:    Straal van de boog in data-eenheden.
    """
    if clockwise:
        theta1, theta2 = -100, 100
    else:
        theta1, theta2 = 80, 280

    arc = Arc(
        (x, y), 2 * radius, 2 * radius,
        angle=0, theta1=theta1, theta2=theta2,
        color=color, linewidth=1.8
    )
    ax.add_patch(arc)

    # Pijlpunt
    angle_rad = math.radians(theta2 if clockwise else theta1)
    tip_x = x + radius * math.cos(angle_rad)
    tip_y = y + radius * math.sin(angle_rad)
    tangent = angle_rad + (math.pi / 2 if clockwise else -math.pi / 2)
    spread = 0.45
    arrow_len = radius * 0.35
    ax.annotate(
        '', xy=(tip_x, tip_y),
        xytext=(
            tip_x - arrow_len * math.cos(tangent),
            tip_y - arrow_len * math.sin(tangent)
        ),
        arrowprops=dict(arrowstyle='->', color=color, lw=1.5),
        clip_on=True,
    )


def draw_zigzag_line(
    ax: Axes,
    x1: float,
    x2: float,
    y: float,
    amplitude: float = 0.12,
    segment: float = 0.2,
    color: str = '#111111',
    linewidth: float = 1.6,
) -> None:
    """Teken een zigzag-lijn (voor veersteun-symbolen).

    Parameters
    ----------
    ax:        Matplotlib Axes.
    x1, x2:   Begin- en eindpunt in data-x.
    y:         Hoogte in data-y.
    amplitude: Uitwijking van de zigzag.
    segment:   Lengte van elk segment.
    color:     Lijnkleur.
    linewidth: Lijndikte.
    """
    length = x2 - x1
    count = max(3, int(abs(length) / segment))
    xs = [x1]
    ys = [y]
    for i in range(1, count + 1):
        t = i / count
        xi = x1 + length * t
        yi = y + (amplitude if i % 2 == 0 else -amplitude)
        xs.append(xi)
        ys.append(yi)
    xs.append(x2)
    ys.append(y)
    ax.plot(xs, ys, color=color, linewidth=linewidth, clip_on=True)

''',

    'tests/test_word_exporter.py': r'''
"""Tests voor WordExporter."""

from __future__ import annotations

import base64
import os
import tempfile

from docx import Document

from exporters import word_exporter
from exporters.word_exporter import WordExporter
from reporting.models import (
    FaseInvoerSectie,
    ReportField,
    ReportImageGroup,
    ReportImageRequest,
    ReportItem,
    ReportMetadata,
    ReportPackage,
    ReportSection,
)
from reporting.builders.input_description_builder import FaseCard, FaseRow


_PNG_1X1 = base64.b64decode(
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwAD'
    'hgGAWjR9awAAAABJRU5ErkJggg=='
)


def test_word_exporter_schrijft_image_als_inline_shape(monkeypatch) -> None:
    """Een sectie met image-request geeft een inline shape in de docx."""
    monkeypatch.setattr(word_exporter, 'render_figuur', lambda _img, _project: _PNG_1X1)
    sec = ReportSection(
        id='fase_1',
        title='Fase 1',
        images=[
            ReportImageRequest(
                id='img_fase_1',
                caption='Doorsnede fase 1',
                figure_key='section',
                stage_index=0,
                step_key=None,
            )
        ],
    )
    pkg = ReportPackage(
        metadata=ReportMetadata(project_name='T'),
        input_sections=[sec],
        selected_items=[
            ReportItem(
                id='damwand_fase_1',
                kind='invoer',
                caption='Fase 1',
                source_ref='fase_1',
            )
        ],
    )
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as handle:
        out = handle.name
    fout = WordExporter().export(pkg, None, out, project=object())
    assert fout is None
    doc = Document(out)
    os.unlink(out)
    assert len(doc.inline_shapes) == 1


def test_word_exporter_schrijft_image_group_als_tabel(monkeypatch) -> None:
    """Een 3x3 figuurgroep wordt als Word-tabel met drie afbeeldingen geschreven."""
    monkeypatch.setattr(word_exporter, 'render_figuur', lambda _img, _project: _PNG_1X1)
    img = ReportImageRequest(
        id='m',
        caption='',
        figure_key='moment_curve',
        stage_index=0,
        step_key='CUR 166 6.4',
    )
    sec = ReportSection(
        id='extremen_overzicht',
        title='Maatgevende resultaten',
        image_groups=[
            ReportImageGroup(
                id='extremen_3x3',
                title='',
                headers=['Msd = 210 kNm/m', 'Dsd = 95 kN/m', 'Urep BGT = 12 mm'],
                images=[img, img, img],
                footers=['Fase 1 - Start', 'Fase 2 - Eind', 'Fase 2 - Eind'],
            )
        ],
    )
    pkg = ReportPackage(
        metadata=ReportMetadata(project_name='T'),
        result_sections=[sec],
        selected_items=[
            ReportItem(
                id='result_extremen_overzicht',
                kind='resultaat',
                caption='Maatgevende resultaten',
                source_ref='extremen_overzicht',
            )
        ],
    )
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as handle:
        out = handle.name
    fout = WordExporter().export(pkg, None, out, project=object())
    assert fout is None
    doc = Document(out)
    os.unlink(out)
    assert len(doc.tables) == 2
    assert len(doc.inline_shapes) == 3
    assert doc.tables[-1].rows[0].cells[0].text == 'Msd = 210 kNm/m'


def test_word_exporter_gebruikt_fase_invoer_tabel_layout() -> None:
    """Normale Word-export gebruikt dezelfde fase-layout als hoofdstukexport."""
    kaart = FaseCard(fase_num=2, stage_name='Fase 2: Belasting')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    kaart.rows.append(FaseRow(
        'Bovenbelasting',
        'op maaiveld',
        '5,0 [kN/mÂ²]',
        extra_lines=['3,0m breed', '0,0m vanaf damwand'],
    ))
    sec = FaseInvoerSectie(
        id='fase_2_invoer',
        title='Fase 2: Belasting',
        fase_card=kaart,
    )
    pkg = ReportPackage(
        metadata=ReportMetadata(project_name='T'),
        input_sections=[sec],
        selected_items=[
            ReportItem(
                id='damwand_fase_2_invoer',
                kind='invoer',
                caption='Fase 2: Belasting',
                source_ref='fase_2_invoer',
            )
        ],
    )
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as handle:
        out = handle.name
    fout = WordExporter().export(pkg, None, out, project=None)
    assert fout is None
    doc = Document(out)
    os.unlink(out)

    assert len(doc.tables) >= 2
    fase_tabel = doc.tables[-1]
    assert [c.text for c in fase_tabel.rows[1].cells][:3] == [
        'Parameter', 'Niveau', 'Toelichting',
    ]
    alle_tekst = '\n'.join(c.text for row in fase_tabel.rows for c in row.cells)
    assert 'Bovenbelasting' in alle_tekst
    assert '3,0m breed' in alle_tekst
    assert '0,0m vanaf damwand' in alle_tekst


def test_word_exporter_gebruikt_damwandgegevens_tabel_layout() -> None:
    """Normale Word-export schrijft damwandgegevens als voorbeeldtabel."""
    sec = ReportSection(id='damwand_gegevens', title='Damwandgegevens')
    sec.fields.extend([
        ReportField('profiel', 'Profiel', 'AZ 18-700'),
        ReportField('hoogte', 'Hoogte', '420,0', 'mm'),
    ])
    pkg = ReportPackage(
        metadata=ReportMetadata(project_name='T'),
        input_sections=[sec],
        selected_items=[
            ReportItem(
                id='damwand_damwand_gegevens',
                kind='invoer',
                caption='Damwandgegevens',
                source_ref='damwand_gegevens',
            )
        ],
    )
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as handle:
        out = handle.name
    fout = WordExporter().export(pkg, None, out, project=None)
    assert fout is None
    doc = Document(out)
    os.unlink(out)

    damwand_tabel = doc.tables[-1]
    assert [c.text for c in damwand_tabel.rows[0].cells] == [
        'Parameter', 'Waarde', 'Eenheid',
    ]
    assert [c.text for c in damwand_tabel.rows[2].cells] == [
        'Hoogte', '420,0', '[mm]',
    ]

''',

    'tests/test_html_preview_builder.py': r'''
"""Tests voor HtmlPreviewBuilder."""

from __future__ import annotations

from reporting.builders.html_preview_builder import HtmlPreviewBuilder
from reporting.builders import html_preview_builder
from reporting.models import (
    ReportImageGroup,
    ReportPackage, ReportMetadata, ReportSection,
    ReportField, ReportImageRequest, ReportTable, ReportItem, TextBlock,
)


def _maak_package_met_sectie() -> tuple[ReportPackage, ReportSection]:
    sec = ReportSection(id='damwand', title='Damwand')
    sec.fields.append(ReportField('top', 'Bovenzijde', '0,00', 'm NAP'))
    sec.fields.append(ReportField('bot', 'Onderzijde', '-12,50', 'm NAP'))
    item = ReportItem(
        id='input_damwand', kind='invoer', caption='Damwand', source_ref='damwand'
    )
    pkg = ReportPackage(
        metadata=ReportMetadata(project_name='Testproject'),
        input_sections=[sec],
        selected_items=[item],
    )
    return pkg, sec


def test_lege_package_bevat_geen_secties_melding() -> None:
    """Lege package â†’ melding 'Geen secties geselecteerd'."""
    html = HtmlPreviewBuilder().build(ReportPackage())
    assert 'Geen secties geselecteerd' in html


def test_projectnaam_zichtbaar_in_html() -> None:
    pkg, _ = _maak_package_met_sectie()
    html = HtmlPreviewBuilder().build(pkg)
    assert 'Testproject' in html


def test_geselecteerde_sectietitel_zichtbaar() -> None:
    pkg, _ = _maak_package_met_sectie()
    html = HtmlPreviewBuilder().build(pkg)
    assert 'Damwand' in html


def test_velden_zichtbaar_in_html() -> None:
    pkg, _ = _maak_package_met_sectie()
    html = HtmlPreviewBuilder().build(pkg)
    assert 'Bovenzijde' in html
    assert '0,00' in html
    assert 'm NAP' in html


def test_niet_geselecteerde_sectie_niet_zichtbaar() -> None:
    sec = ReportSection(id='water', title='Waterpeilen')
    pkg = ReportPackage(input_sections=[sec], selected_items=[])
    html = HtmlPreviewBuilder().build(pkg)
    assert 'Waterpeilen' not in html


def test_tabel_weergegeven_als_html_tabel() -> None:
    tbl = ReportTable(
        id='ankers', title='Ankers',
        columns=['Naam', 'Niveau [m NAP]'],
        rows=[['Anker-1', '-3,50'], ['Anker-2', '-4,00']],
    )
    sec = ReportSection(id='anchors', title='Ankers', tables=[tbl])
    item = ReportItem(id='input_anchors', kind='invoer', caption='Ankers', source_ref='anchors')
    pkg = ReportPackage(input_sections=[sec], selected_items=[item])
    html = HtmlPreviewBuilder().build(pkg)
    assert '<table' in html.lower()
    assert 'Anker-1' in html
    assert '-3,50' in html


def test_tekstblok_effective_text_zichtbaar() -> None:
    sec = ReportSection(id='sec', title='Sectie')
    sec.text_blocks.append(TextBlock(
        id='blk', section='sec',
        generated_text='Gegenereerde tekst.',
        manual_override='Handmatige override.',
    ))
    item = ReportItem(id='input_sec', kind='invoer', caption='Sectie', source_ref='sec')
    pkg = ReportPackage(input_sections=[sec], selected_items=[item])
    html = HtmlPreviewBuilder().build(pkg)
    assert 'Handmatige override.' in html
    assert 'Gegenereerde tekst.' not in html


def test_resultaat_sectie_opgenomen_bij_kind_resultaat() -> None:
    sec = ReportSection(id='moment', title='Momentendiagram')
    sec.fields.append(ReportField('max', 'Max. moment', '142,3', 'kNm/m'))
    item = ReportItem(
        id='result_moment', kind='resultaat', caption='Momenten', source_ref='moment'
    )
    pkg = ReportPackage(result_sections=[sec], selected_items=[item])
    html = HtmlPreviewBuilder().build(pkg)
    assert 'Momentendiagram' in html
    assert '142,3' in html


def test_grondsoorten_sectie_opgenomen_bij_kind_grondsoorten() -> None:
    """kind='grondsoorten' â†’ sectie uit extra_sections zichtbaar in HTML."""
    sec = ReportSection(id='soil_table_links', title='Grondsoortentabel \u2014 Links')
    tbl = ReportTable(
        id='t', title='',
        columns=['BK [m NAP]', 'Laag'],
        rows=[['-5,0', 'Zand']],
    )
    sec.tables.append(tbl)
    item = ReportItem(
        id='grondsoorten_soil_table_links',
        kind='grondsoorten',
        caption='Grondsoortentabel \u2014 Links',
        source_ref='soil_table_links',
    )
    pkg = ReportPackage(extra_sections=[sec], selected_items=[item])
    html = HtmlPreviewBuilder().build(pkg)
    assert 'Grondsoortentabel' in html
    assert 'Zand' in html


def test_image_request_wordt_data_uri_met_caption(monkeypatch) -> None:
    """Figuurverzoeken worden als base64-afbeelding in HTML opgenomen."""
    monkeypatch.setattr(html_preview_builder, 'render_figuur', lambda _img, _project: b'png')
    sec = ReportSection(
        id='fase_1',
        title='Fase 1',
        images=[
            ReportImageRequest(
                id='fig_1',
                caption='Doorsnede fase 1',
                figure_key='section',
                stage_index=0,
                step_key=None,
            )
        ],
    )
    item = ReportItem(
        id='damwand_fase_1',
        kind='invoer',
        caption='Fase 1',
        source_ref='fase_1',
    )
    pkg = ReportPackage(input_sections=[sec], selected_items=[item])
    html = HtmlPreviewBuilder().build(pkg, project=object())
    assert 'data:image/png;base64,cG5n' in html
    assert 'Doorsnede fase 1' in html


def test_image_group_wordt_als_3x3_tabel_gerenderd(monkeypatch) -> None:
    """Een figuurgroep rendert headers, figuurcellen en bronregels in HTML."""
    monkeypatch.setattr(html_preview_builder, 'render_figuur', lambda _img, _project: b'png')
    img = ReportImageRequest(
        id='m',
        caption='',
        figure_key='moment_curve',
        stage_index=0,
        step_key='CUR 166 6.4',
    )
    sec = ReportSection(
        id='extremen_overzicht',
        title='Maatgevende resultaten',
        image_groups=[
            ReportImageGroup(
                id='extremen_3x3',
                title='',
                headers=['Msd = 210 kNm/m', 'Dsd = 95 kN/m', 'Urep BGT = 12 mm'],
                images=[img, img, img],
                footers=['Fase 1 - Start', 'Fase 2 - Eind', 'Fase 2 - Eind'],
            )
        ],
    )
    item = ReportItem(
        id='result_extremen_overzicht',
        kind='resultaat',
        caption='Maatgevende resultaten',
        source_ref='extremen_overzicht',
    )
    pkg = ReportPackage(result_sections=[sec], selected_items=[item])

    html = HtmlPreviewBuilder().build(pkg, project=object())

    assert 'class="figuurgroep"' in html
    assert 'Msd = 210 kNm/m' in html
    assert html.count('data:image/png;base64,cG5n') == 3
    assert 'Fase 1 - Start' in html

''',

}


# ---------------------------------------------------------------------------
# 2026-05-11 - Verwijderde .shi/.shs-runtime parsing
# ---------------------------------------------------------------------------

DEAD_SHI_SHS_RUNTIME_PARSING = {
    "reden": (
        "De app gebruikt .shd als bron van waarheid. Controle op 62 .shi/.shd-paren "
        "in de testfile-map liet zien dat de parseruitkomsten voor invoerdata gelijk "
        "zijn, terwijl .shi/.shs combineren dubbele [SOIL]-blokken kan veroorzaken."
    ),
    "parsers.models.FileBundle": r'''
@dataclass
class FileBundle:
    shi: str = ""
    shd: str = ""
    shs: str = ""
''',
    "parsers.shi_parser.parse_project_oude_bronselectie": r'''
shi = file_bundle.shi or ''
shd = file_bundle.shd or ''
shs = file_bundle.shs or ''
source = shi or shd or shs
combined = '\n'.join(t for t in [shi, shd, shs] if t)

soils = parse_soils(combined)
profiles = parse_soil_profiles(combined)
surfaces = parse_surfaces(combined)
waterlevels = parse_water_levels(combined)
sheet_piling = parse_sheet_piling(combined)
anchors = parse_anchors(combined)
struts = parse_struts(combined)
spring_supports = parse_spring_supports(combined)
rigid_supports = parse_rigid_supports(combined)
uniform_loads = parse_uniform_loads(combined)
surcharge_loads = parse_surcharge_loads(combined)
horizontal_line_loads = parse_horizontal_line_loads(combined)
moments = parse_moments(combined)
normal_forces = parse_normal_forces(combined)
stages = parse_stages(combined)
''',
    "app.controller.ingest_process_remove_group_oude_extensies": r'''
if ext not in ('shi', 'shd', 'shs'):
    continue

grouped[base] = {'shi': '', 'shd': '', 'shs': ''}
if ext in ('shi', 'shd', 'shs'):
    grouped[base][ext] = text

for ext in ('shi', 'shd', 'shs'):
    self._state.raw_files.pop(f'{base_name}.{ext}', None)

return re.sub(r'\.(shi|shd|shs)$', '', filename, flags=re.IGNORECASE)
''',
    "ui_oude_bestandsindicatie_en_filter": r'''
('D-Sheet bestanden (*.shi *.shd *.shs);;Alle bestanden (*)')

('.shi',  'ja' if project.file_bundle.shi else 'nee')
('.shd',  'ja' if project.file_bundle.shd else 'nee')
('.shs',  'ja' if project.file_bundle.shs else 'nee')
''',
}
