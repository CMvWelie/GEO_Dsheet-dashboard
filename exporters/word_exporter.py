"""WordExporter — exporteert een ReportPackage naar Word (.docx).

JSON-sidecar formaat (naast .dotx sjabloon, zelfde naam + .map.json):
{
  "metadata": {
    "project_name": "bookmark_project",
    "title":        "bookmark_title"
  },
  "sections": {
    "sheet_piling": "Sectie 2.1 Damwand",
    "moment_max":   "Sectie 3.1 Momenten"
  }
}
Sleutels in 'metadata' zijn bladwijzernamen in het template.
Sleutels in 'sections' zijn koppen waaronder de secties worden ingevoegd.
"""

from __future__ import annotations
import io
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

from reporting.models import FaseInvoerSectie, ReportPackage, ReportSection
from reporting.figure_renderer import render_figuur

_DOTX_CONTENT_TYPE = (
    'application/vnd.openxmlformats-officedocument'
    '.wordprocessingml.template.main+xml'
)
_DOCX_CONTENT_TYPE = (
    'application/vnd.openxmlformats-officedocument'
    '.wordprocessingml.document.main+xml'
)


class WordExporter:
    """Schrijft een ReportPackage naar een .docx-bestand met python-docx."""

    def export(self, package: ReportPackage, template_path: str | None,
               output_path: str, project=None) -> str | None:
        """Exporteer naar Word.

        Returns:
            None bij succes, foutmelding (str) bij een uitzondering.
        """
        try:
            mapping = self._load_mapping(template_path)

            if template_path and Path(template_path).exists():
                doc = self._open_template(template_path)
            else:
                doc = Document()
            self._apply_theme_typography(doc)

            if mapping:
                self._write_with_mapping(doc, package, mapping, project)
            else:
                self._write_metadata(doc, package)
                selected_ids = {
                    i.source_ref for i in package.selected_items
                    if i.source_ref and i.included_word
                }
                # extra_sections worden alleen gefilterd als er een overeenkomend
                # ReportItem met source_ref in selected_items bestaat.
                all_sections = (
                    package.input_sections
                    + package.result_sections
                    + package.extra_sections
                )
                for sec in all_sections:
                    if selected_ids and sec.id not in selected_ids:
                        continue
                    self._write_section(doc, sec, project)

            doc.save(output_path)
            return None
        except Exception as exc:
            return str(exc)

    # ------------------------------------------------------------------
    # Sjabloon openen (.dotx én .docx)
    # ------------------------------------------------------------------

    def _open_template(self, path: str) -> Document:
        """Open een Word-sjabloon als bewerkbaar Document.

        Een .dotx-bestand heeft een ander content type dan .docx, waardoor
        python-docx het weigert. We passen het content type in-memory aan
        zodat Document() het accepteert — zonder tijdelijke bestanden.

        Parameters
        ----------
        path: Pad naar het sjabloonbestand (.dotx of .docx).
        """
        if Path(path).suffix.lower() != '.dotx':
            return Document(path)

        with open(path, 'rb') as f:
            data = f.read()

        invoer = io.BytesIO(data)
        uitvoer = io.BytesIO()
        with zipfile.ZipFile(invoer, 'r') as zin, \
                zipfile.ZipFile(uitvoer, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                inhoud = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    inhoud = inhoud.replace(
                        _DOTX_CONTENT_TYPE.encode(),
                        _DOCX_CONTENT_TYPE.encode(),
                    )
                zout.writestr(item, inhoud)

        uitvoer.seek(0)
        return Document(uitvoer)

    def _apply_theme_typography(self, doc) -> None:
        """Pas thematekstgrootte toe op standaardtekst buiten tabellen."""
        from exporters.word_hoofdstuk_exporter import _eerste_fontfamilie
        from ui import table_styles

        stijl = doc.styles['Normal']
        stijl.font.name = _eerste_fontfamilie(table_styles.TABLE_FONT)
        stijl.font.size = Pt(table_styles.WORD_BODY_TEXT_SIZE)

    # ------------------------------------------------------------------
    # JSON-sidecar laden
    # ------------------------------------------------------------------

    def _load_mapping(self, template_path: str | None) -> dict | None:
        if not template_path:
            return None
        sidecar = Path(template_path).with_suffix('').with_suffix('.map.json')
        if not sidecar.exists():
            sidecar2 = Path(str(template_path) + '.map.json')
            if not sidecar2.exists():
                return None
            sidecar = sidecar2
        try:
            return json.loads(sidecar.read_text(encoding='utf-8'))
        except Exception:
            return None

    # ------------------------------------------------------------------
    # Schrijven met sidecar-mapping
    # ------------------------------------------------------------------

    def _write_with_mapping(
        self, doc, package: ReportPackage, mapping: dict, project=None
    ) -> None:
        """Vul template-bladwijzers in en voeg secties toe na genoemde koppen."""
        # Metadata via bladwijzers
        meta_map = mapping.get('metadata', {})
        md = package.metadata
        for attr, bookmark_name in meta_map.items():
            value = getattr(md, attr, '') or ''
            self._fill_bookmark(doc, bookmark_name, value)

        # Secties
        sec_map = mapping.get('sections', {})
        all_sections = (
            package.input_sections
            + package.result_sections
            + package.extra_sections
        )
        for sec in all_sections:
            heading_text = sec_map.get(sec.id)
            if heading_text:
                self._insert_after_heading(doc, heading_text, sec, project)
            else:
                self._write_section(doc, sec, project)

    def _fill_bookmark(self, doc, bookmark_name: str, value: str) -> None:
        """Vervang tekst van een bladwijzer in het document."""
        from docx.oxml.ns import qn
        for para in doc.paragraphs:
            for bm in para._element.findall(f'.//{qn("w:bookmarkStart")}'):
                if bm.get(qn('w:name')) == bookmark_name:
                    # Verwijder bestaande runs in de paragraaf en voeg waarde toe
                    for run in para.runs:
                        run.text = ''
                    if para.runs:
                        para.runs[0].text = value
                    else:
                        para.add_run(value)
                    return

    def _insert_after_heading(self, doc, heading_text: str,
                               section: ReportSection, project=None) -> None:
        """Voeg sectie-content toe na de paragraaf met de gegeven koptekst."""
        from docx.oxml import OxmlElement
        target = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == heading_text.strip():
                target = para._element
                break
        if target is None:
            # Kop niet gevonden → aan het einde toevoegen
            self._write_section(doc, section)
            return
        # Bouw sectie-XML op en voeg in na de kop
        temp_doc = doc.__class__()
        self._write_section(temp_doc, section, project)
        parent = target.getparent()
        idx = list(parent).index(target) + 1
        for elem in list(temp_doc.element.body)[:-1]:  # skip sectPr
            parent.insert(idx, elem)
            idx += 1

    # ------------------------------------------------------------------
    # Standaard schrijven (geen mapping)
    # ------------------------------------------------------------------

    def _write_metadata(self, doc, package: ReportPackage) -> None:
        md = package.metadata
        doc.add_heading('Rapportgegevens', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = 'Veld'
        table.rows[0].cells[1].text = 'Waarde'
        rows = [
            ('Projectnaam',   md.project_name),
            ('Ordernummer',   md.order_number),
            ('Locatie',       md.location),
            ('Fase',          md.phase),
            ('Opdrachtgever', md.client),
            ('Titel',         md.title),
            ('Revisie',       md.revision),
            ('Auteur',        md.author),
            ('Datum',         md.date),
        ]
        for label, value in rows:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value or ''

    def _write_section(self, doc, section: ReportSection, project=None) -> None:
        if isinstance(section, FaseInvoerSectie):
            from exporters.word_hoofdstuk_exporter import WordHoofdstukExporter

            WordHoofdstukExporter()._schrijf_fase_sectie(doc, section, project)
            return
        if section.id == 'damwand_gegevens':
            from exporters.word_hoofdstuk_exporter import WordHoofdstukExporter

            WordHoofdstukExporter()._schrijf_damwandgegevens_sectie(doc, section)
            return

        doc.add_heading(section.title, level=2)

        for f in section.fields:
            val = f'{f.value} {f.unit}'.strip() if f.unit else f.value
            doc.add_paragraph(f'{f.label}: {val}')

        for table_data in section.tables:
            doc.add_paragraph(table_data.title, style='Intense Quote')
            if not table_data.columns:
                continue
            tbl = doc.add_table(rows=1, cols=len(table_data.columns))
            for col, header in enumerate(table_data.columns):
                tbl.rows[0].cells[col].text = header
            for data_row in table_data.rows:
                row = tbl.add_row()
                for col, cell in enumerate(data_row):
                    if col < len(row.cells):
                        row.cells[col].text = cell

        for tb in section.text_blocks:
            doc.add_paragraph(tb.effective_text)

        for groep in section.image_groups:
            self._write_image_group(doc, groep, project)

        for img_req in section.images:
            self._write_image(doc, img_req, project)

    def _write_image(self, doc, img_req, project) -> None:
        """Render een figuurverzoek en voeg het toe aan het Word-document."""
        if project is None:
            doc.add_paragraph(f'[Figuur: {img_req.caption}]')
            return
        png_bytes = render_figuur(img_req, project)
        if png_bytes:
            doc.add_picture(io.BytesIO(png_bytes), width=Cm(16))
            if img_req.caption:
                doc.add_paragraph(img_req.caption, style='Caption')
        else:
            doc.add_paragraph(f'[Figuur niet beschikbaar: {img_req.caption}]')

    def _write_image_group(self, doc, groep, project) -> None:
        """Schrijf een figuurgroep als Word-tabel met kop, figuur en bron."""
        if not groep.headers:
            return
        if groep.title:
            doc.add_paragraph(groep.title, style='Intense Quote')

        tbl = doc.add_table(rows=3, cols=len(groep.headers))
        for col, header in enumerate(groep.headers):
            tbl.rows[0].cells[col].text = header

        for col, img_req in enumerate(groep.images):
            cell = tbl.rows[1].cells[col]
            if img_req is None:
                cell.text = '-'
                continue
            if project is None:
                cell.text = f'[Figuur: {img_req.caption or img_req.figure_key}]'
                continue
            png_bytes = render_figuur(img_req, project)
            if not png_bytes:
                cell.text = '[Figuur niet beschikbaar]'
                continue
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(io.BytesIO(png_bytes), width=Cm(5.2))

        for col, footer in enumerate(groep.footers):
            if col < len(tbl.rows[2].cells):
                tbl.rows[2].cells[col].text = footer
