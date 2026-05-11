"""WordHoofdstukExporter — exporteert het damwand-hoofdstuk naar Word."""
from __future__ import annotations
import io
import zipfile
from pathlib import Path
import re
import struct

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table

from reporting.models import (
    FaseInvoerSectie,
    ReportField,
    ReportImageGroup,
    ReportImageRequest,
    ReportMetadata,
    ReportSection,
)
from reporting.builders.damwand_tekst import (
    FASERING_INTRO_TEKST,
    FASERING_TABEL_INTRO_TEKST,
    FASERING_TITEL,
    RESULTATEN_GRAFIEK_INTRO_TEKST,
    RESULTATEN_INTRO_TEKST,
    RESULTATEN_TITEL,
    faseringsregels,
    project_fase_namen,
)
from reporting.figure_renderer import render_figuur


_FASE_RIJHOOGTE_CM = 0.45
_DAMWAND_KOLOM_BREEDTES_CM = [5.0, 3.0, 2.0]
_RESULTAAT_SPEC_KOLOM_BREEDTES_CM = [5.0, 2.5, 3.5, 2.0]  # label, stap, waarde, eenheid
_GRONDSOORTEN_V2_FASE_KOLOM_BREEDTES_CM = [4.0, 2.0, 2.0, 4.0, 2.0, 2.0]
_DAMWAND_RIJHOOGTE_TWIPS = round(_FASE_RIJHOOGTE_CM * 567)
_DAMWAND_SCHEIDING_TWIPS = 40
_RESULTAAT_SECTIE_IDS = {
    'conclusietabel',
    'anchor_forces',
    'per_phase_summary',
    'extremen_overzicht',
    'grafieken_moment_dwarskracht',
    'grafieken_vervorming',
}


def _png_hoogte_cm(png_bytes: bytes, breedte_cm: float) -> float:
    """Bereken de weergavehoogte in cm van een PNG bij vaste breedte.

    Parameters
    ----------
    png_bytes:
        PNG-bestand als bytes.
    breedte_cm:
        Gewenste weergavebreedte in cm.

    Returns
    -------
    float
        De bijbehorende hoogte in cm, of 0,0 bij ongeldige invoer.
    """
    if len(png_bytes) < 24:
        return 0.0
    w_px = struct.unpack('>I', png_bytes[16:20])[0]
    h_px = struct.unpack('>I', png_bytes[20:24])[0]
    if w_px == 0:
        return 0.0
    return breedte_cm * h_px / w_px


def _hex_zonder_hash(kleur: str) -> str:
    """Normaliseer een CSS-hexkleur naar Word-hex zonder ``#``."""
    kleur = (kleur or '').strip()
    if kleur.startswith('#'):
        kleur = kleur[1:]
    return kleur.upper() if re.fullmatch(r'[0-9A-Fa-f]{6}', kleur) else '000000'


def _eerste_fontfamilie(font_stack: str) -> str:
    """Haal de eerste fontnaam uit een CSS-font-stack."""
    eerste = (font_stack or 'Arial').split(',')[0].strip()
    return eerste.strip('"\'') or 'Arial'


class WordHoofdstukExporter:
    """Schrijft een lijst van ReportSection objecten naar een .docx bestand."""

    def export(
        self,
        sections: list[ReportSection],
        metadata: ReportMetadata,
        project: object | None,
        template_path: str | None,
        output_path: str,
    ) -> str | None:
        """Exporteer naar Word.

        Parameters
        ----------
        sections:      Gesorteerde lijst van ReportSection objecten.
        metadata:      Rapportagegegevens (projectnaam, auteur, etc.).
        project:       Project-object voor figuurrendering (mag None zijn als geen figuren).
        template_path: Pad naar het .docx stijlenbestand, of None voor leeg document.
        output_path:   Uitvoerpad voor het .docx bestand.

        Returns
        -------
        str | None
            None bij succes, foutmelding bij uitzondering.
        """
        try:
            doc = self._open_doc(template_path)
            self._pas_document_typografie_toe(doc)
            self._schrijf_titel(doc, metadata)
            resultaat_specificaties_geschreven = False
            faserings_intro_geschreven = False
            self._fase_tabel_teller = 0
            fase_secties = [
                sec for sec in sections if isinstance(sec, FaseInvoerSectie)
            ]
            for sec in sections:
                if (
                    not resultaat_specificaties_geschreven
                    and project is not None
                    and sec.id in _RESULTAAT_SECTIE_IDS
                ):
                    self._schrijf_resultaat_specificaties_tabel(doc, project)
                    resultaat_specificaties_geschreven = True
                if (
                    isinstance(sec, FaseInvoerSectie)
                    and not faserings_intro_geschreven
                ):
                    self._schrijf_fasering_intro(doc, project, fase_secties)
                    faserings_intro_geschreven = True
                self._schrijf_sectie(doc, sec, project)
            doc.save(output_path)
            return None
        except Exception as exc:
            return str(exc)

    # ------------------------------------------------------------------
    # Document openen
    # ------------------------------------------------------------------

    def _open_doc(self, template_path: str | None) -> Document:
        if not template_path or not Path(template_path).exists():
            return Document()
        # .dotx heeft ander content-type dan python-docx verwacht; patch naar .docx
        if Path(template_path).suffix.lower() == '.dotx':
            buf = io.BytesIO()
            with zipfile.ZipFile(template_path, 'r') as zin, \
                 zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == '[Content_Types].xml':
                        data = data.replace(
                            b'wordprocessingml.template.main+xml',
                            b'wordprocessingml.document.main+xml',
                        )
                    zout.writestr(item, data)
            buf.seek(0)
            return Document(buf)
        return Document(template_path)

    # ------------------------------------------------------------------
    # Titel
    # ------------------------------------------------------------------

    def _schrijf_titel(self, doc: Document, metadata: ReportMetadata) -> None:
        titel = metadata.project_name or 'Damwand rapport'
        doc.add_heading(titel, level=1)

    def _pas_document_typografie_toe(self, doc: Document) -> None:
        """Pas thematekstgrootte toe op standaardtekst buiten tabellen."""
        from ui import table_styles

        stijl = doc.styles['Normal']
        stijl.font.name = _eerste_fontfamilie(table_styles.TABLE_FONT)
        stijl.font.size = Pt(table_styles.WORD_BODY_TEXT_SIZE)

    # ------------------------------------------------------------------
    # Secties
    # ------------------------------------------------------------------

    def _schrijf_sectie(
        self,
        doc: Document,
        sec: ReportSection,
        project: object | None,
    ) -> None:
        if isinstance(sec, FaseInvoerSectie):
            self._schrijf_fase_sectie(doc, sec, project)
            return
        if sec.id == 'damwand_gegevens':
            self._schrijf_damwandgegevens_sectie(doc, sec)
            return
        if sec.id in {'per_phase_summary', 'anchor_forces'}:
            return
        if sec.id == 'extremen_overzicht':
            self._schrijf_extremen_overzicht_sectie(doc, sec, project)
            return
        if sec.id.startswith('grondsoorten_v2_fase_'):
            self._schrijf_grondsoorten_v2_fase_sectie(doc, sec)
            return

        doc.add_heading(sec.title, level=2)
        for veld in sec.fields:
            waarde = f'{veld.value} {veld.unit}'.strip() if veld.unit else veld.value
            doc.add_paragraph(f'{veld.label}: {waarde}')
        for tabel in sec.tables:
            self._schrijf_tabel(doc, tabel)
        for tb in sec.text_blocks:
            doc.add_paragraph(tb.effective_text)
        for img_req in sec.images:
            self._schrijf_figuur(doc, img_req, project)
        for groep in sec.image_groups:
            self._schrijf_figuurgroep(doc, groep, project)

    def _schrijf_grondsoorten_v2_fase_sectie(
        self,
        doc: Document,
        sec: ReportSection,
    ) -> None:
        """Schrijf een Grondsoortentabel v2-fasesectie met intro voor de tabel."""
        if sec.title:
            doc.add_heading(sec.title, level=2)
        else:
            doc.add_paragraph()
        for tb in sec.text_blocks:
            tekst = tb.effective_text
            self._schrijf_toelichting_met_bullets(doc, tekst)
            if len((tekst or '').splitlines()) > 1 or not sec.title:
                doc.add_paragraph()
        for tabel in sec.tables:
            self._schrijf_tabel(doc, tabel)

    def _schrijf_damwandgegevens_sectie(
        self,
        doc: Document,
        sec: ReportSection,
    ) -> None:
        """Schrijf damwandgegevens als compacte 3-koloms Word-tabel."""
        kop = doc.add_heading(sec.title, level=2)
        kop.paragraph_format.page_break_before = True
        for tb in sec.text_blocks[:1]:
            self._schrijf_textblock(doc, tb.effective_text)
        doc.add_paragraph()
        if not sec.fields:
            for tb in sec.text_blocks[1:]:
                self._schrijf_toelichting_met_bullets(doc, tb.effective_text)
            return

        velden: list[ReportField | None] = []
        for veld in sec.fields:
            if veld.key.startswith('ondersteuning_') and not any(v is None for v in velden):
                velden.append(None)
            velden.append(veld)

        tbl = doc.add_table(rows=1 + len(velden), cols=3)
        tbl.autofit = False
        try:
            tbl.style = 'Table Grid'
        except KeyError:
            pass

        kolom_breedtes = [round(cm * 567) for cm in _DAMWAND_KOLOM_BREEDTES_CM]
        self._stel_tabel_grid_in(tbl, _DAMWAND_KOLOM_BREEDTES_CM)
        for row in tbl.rows:
            for col_idx, cell in enumerate(row.cells[:3]):
                self._stel_cel_breedte(cell, kolom_breedtes[col_idx])
            self._stel_rijhoogte_exact_twips(row, _DAMWAND_RIJHOOGTE_TWIPS)

        for col, tekst in enumerate(['Parameter', 'Waarde', 'Eenheid']):
            tbl.rows[0].cells[col].text = tekst
        tbl.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row_idx, veld in enumerate(velden, start=1):
            if veld is None:
                self._stel_rijhoogte_exact_twips(
                    tbl.rows[row_idx],
                    _DAMWAND_SCHEIDING_TWIPS,
                )
                continue
            tbl.rows[row_idx].cells[0].text = veld.label
            tbl.rows[row_idx].cells[1].text = veld.value
            tbl.rows[row_idx].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            tbl.rows[row_idx].cells[2].text = self._formatteer_eenheid(veld.unit)

        self._pas_damwand_tabel_opmaak_toe(tbl)
        doc.add_paragraph()
        for tb in sec.text_blocks[1:]:
            self._schrijf_toelichting_met_bullets(doc, tb.effective_text)

    def _schrijf_textblock(self, doc: Document, tekst: str) -> None:
        """Schrijf een tekstblok als een of meer Word-paragrafen."""
        for regel in (tekst or '').splitlines():
            doc.add_paragraph(regel)

    def _schrijf_toelichting_met_bullets(self, doc: Document, tekst: str) -> None:
        """Schrijf eerste regel normaal, overige regels als streepjes-bullets."""
        regels = (tekst or '').splitlines()
        for i, regel in enumerate(regels):
            if i == 0:
                doc.add_paragraph(regel)
            else:
                self._voeg_bullet_paragraaf_toe(doc, regel)

    def _maak_streepjes_bullet_num_id(self, doc: Document) -> str:
        """Voeg een abstractNum + num toe met '-' als bullet; geeft numId terug."""
        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element

        bestaande_abstract = numbering_elm.findall(qn('w:abstractNum'))
        nieuw_abstract_id = str(len(bestaande_abstract))

        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(qn('w:abstractNumId'), nieuw_abstract_id)
        multi = OxmlElement('w:multiLevelType')
        multi.set(qn('w:val'), 'singleLevel')
        abstract_num.append(multi)

        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')
        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'bullet')
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), '-')
        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        p_pr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')
        ind.set(qn('w:hanging'), '360')
        p_pr.append(ind)
        lvl.append(num_fmt)
        lvl.append(lvl_text)
        lvl.append(lvl_jc)
        lvl.append(p_pr)
        abstract_num.append(lvl)
        numbering_elm.append(abstract_num)

        bestaande_num = numbering_elm.findall(qn('w:num'))
        nieuw_num_id = str(len(bestaande_num) + 1)
        num_elm = OxmlElement('w:num')
        num_elm.set(qn('w:numId'), nieuw_num_id)
        abstract_num_id_ref = OxmlElement('w:abstractNumId')
        abstract_num_id_ref.set(qn('w:val'), nieuw_abstract_id)
        num_elm.append(abstract_num_id_ref)
        numbering_elm.append(num_elm)

        return nieuw_num_id

    def _voeg_bullet_paragraaf_toe(self, doc: Document, tekst: str) -> None:
        """Voeg een paragraaf toe met '-' als Word-opsommingsteken."""
        if not hasattr(self, '_bullet_num_id'):
            self._bullet_num_id = self._maak_streepjes_bullet_num_id(doc)
        para = doc.add_paragraph(tekst, style='List Paragraph')
        pPr = para._element.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numId_elm = OxmlElement('w:numId')
        numId_elm.set(qn('w:val'), self._bullet_num_id)
        numPr.append(ilvl)
        numPr.append(numId_elm)
        pPr.append(numPr)

    def _schrijf_fasering_intro(
        self,
        doc: Document,
        project: object | None,
        fase_secties: list[FaseInvoerSectie],
    ) -> None:
        """Schrijf de gezamenlijke inleiding voor alle faseringstabellen."""
        kop = doc.add_heading(FASERING_TITEL, level=2)
        kop.paragraph_format.page_break_before = True
        doc.add_paragraph(FASERING_INTRO_TEKST)
        fase_namen = project_fase_namen(project) or [
            sec.fase_card.stage_name
            for sec in fase_secties
            if sec.fase_card is not None
        ]
        for regel in faseringsregels(fase_namen):
            self._voeg_bullet_paragraaf_toe(doc, regel)
        doc.add_paragraph()
        doc.add_paragraph(FASERING_TABEL_INTRO_TEKST)
        doc.add_paragraph()

    def _formatteer_eenheid(self, eenheid: str) -> str:
        """Geef een Word-eenheid terug met blokhaken, indien aanwezig."""
        tekst = (eenheid or '').strip()
        if not tekst:
            return ''
        if tekst.startswith('[') and tekst.endswith(']'):
            return tekst
        return f'[{tekst}]'

    def _pas_damwand_tabel_opmaak_toe(self, tbl: Table) -> None:
        """Pas themakleuren en fontgroottes toe op de damwandgegevens-tabel."""
        from ui import table_styles

        font = _eerste_fontfamilie(table_styles.TABLE_FONT)
        for row in tbl.rows:
            for cell in row.cells:
                self._pas_cel_font_toe(
                    cell, font, '000000', bold=False,
                    size_pt=table_styles.WORD_TABLE_TEXT_SIZE,
                )

        header_bg = _hex_zonder_hash(table_styles.TABLE_HEADER_BG)
        header_fg = _hex_zonder_hash(table_styles.TABLE_HEADER_FG)
        for cell in tbl.rows[0].cells:
            self._stel_cel_vulling(cell, header_bg)
            self._pas_cel_font_toe(
                cell, font, header_fg, bold=True,
                size_pt=table_styles.WORD_TABLE_HEADER_SIZE,
            )

    def _schrijf_resultaat_specificaties_tabel(
        self,
        doc: Document,
        project: object,
    ) -> None:
        """Schrijf de resultaatspecificatietabel: 4 kolommen [label|stap|waarde|eenheid].

        Rijen zonder stap krijgen col 0+1 samengevoegd (ziet er uit als 3 kolommen).
        """
        rijen = self._resultaat_specificatie_rijen(project)
        if not rijen:
            return

        kop = doc.add_heading(RESULTATEN_TITEL, level=2)
        kop.paragraph_format.page_break_before = True
        doc.add_paragraph(RESULTATEN_INTRO_TEKST)
        doc.add_paragraph()

        kolom_breedtes = [round(cm * 567) for cm in _RESULTAAT_SPEC_KOLOM_BREEDTES_CM]
        label_b, stap_b, waarde_b, eenheid_b = kolom_breedtes
        merged_label_b = label_b + stap_b

        tbl = doc.add_table(rows=len(rijen), cols=4)
        tbl.autofit = False
        try:
            tbl.style = 'Table Grid'
        except KeyError:
            pass
        self._stel_tabel_grid_in(tbl, _RESULTAAT_SPEC_KOLOM_BREEDTES_CM)
        for row in tbl.rows:
            for col_idx, cell in enumerate(row.cells[:4]):
                self._stel_cel_breedte(cell, kolom_breedtes[col_idx])
            self._stel_rijhoogte_exact_twips(row, _DAMWAND_RIJHOOGTE_TWIPS)

        for row_idx, (label, waarde, eenheid, stap) in enumerate(rijen):
            row = tbl.rows[row_idx]
            is_koprij = waarde == '' and eenheid == ''

            if is_koprij and not stap:
                # "Grondkering"-type: alle 4 kolommen samenvoegen
                merged = row.cells[0].merge(row.cells[3])
                self._stel_cel_breedte(merged, sum(kolom_breedtes))
                merged.text = label

            elif is_koprij:
                # "Resultaten"-type: [label | stap-header | rest leeg]
                row.cells[0].text = label
                row.cells[1].text = stap
                merged_rest = row.cells[2].merge(row.cells[3])
                self._stel_cel_breedte(merged_rest, waarde_b + eenheid_b)

            elif stap:
                # Datarij mét stap: [label | stap | waarde | eenheid]
                row.cells[0].text = label
                row.cells[1].text = stap
                row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[2].text = waarde
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[3].text = eenheid

            else:
                # Datarij zonder stap: col 0+1 samenvoegen → [merged-label | waarde | eenheid]
                waarde_cel = row.cells[2]
                eenheid_cel = row.cells[3]
                merged_label = row.cells[0].merge(row.cells[1])
                self._stel_cel_breedte(merged_label, merged_label_b)
                merged_label.text = label
                waarde_cel.text = waarde
                waarde_cel.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                eenheid_cel.text = eenheid

        for row in tbl.rows:
            self._stel_rijhoogte_exact_twips(row, _DAMWAND_RIJHOOGTE_TWIPS)

        self._pas_resultaat_specificaties_opmaak_toe(tbl)

        # Verticale celsamenvoeging voor opeenvolgende rijen met gelijk label
        for j in range(len(rijen) - 1):
            sec_j = rijen[j][1] == '' and rijen[j][2] == ''
            sec_j1 = rijen[j + 1][1] == '' and rijen[j + 1][2] == ''
            if not sec_j and not sec_j1 and rijen[j][0] == rijen[j + 1][0]:
                self._stel_verticale_celmerge_in(
                    tbl.rows[j].cells[0], tbl.rows[j + 1].cells[0]
                )

        doc.add_paragraph()

    def _resultaat_specificatie_rijen(
        self,
        project: object,
    ) -> list[tuple[str, str, str, str]]:
        """Bouw de rijen voor de projectbrede resultaat-specificatietabel."""
        from reporting.builders.result_description_builder import (
            _step_short_label,
            _VTYPE_LABELS,
            is_bgt_step_key,
            is_ugt_step_key,
        )
        from utils.formatting import fmt_number

        sheet_piling = getattr(project, 'sheet_piling', None) or []
        summaries = getattr(project, 'result_summaries', None) or []
        anchor_strut_resume = getattr(project, 'anchor_strut_resume', None) or []
        supports_resume = getattr(project, 'supports_resume', None) or []
        spring_supports = getattr(project, 'spring_supports', None) or []
        rigid_supports = getattr(project, 'rigid_supports', None) or []
        el = sheet_piling[0] if sheet_piling else None

        max_mob_mom = max((s.mob_moment_pct for s in summaries), default=0.0)
        max_mob_grond = max((s.mob_grond_pct for s in summaries), default=0.0)

        msd, dsd, vervorming = self._maatgevende_resultaatwaarden_en_stappen(
            project,
            is_ugt_step_key,
            is_bgt_step_key,
        )

        top = getattr(el, 'top', None) if el is not None else None
        bottom = getattr(el, 'bottom', 0.0) if el is not None else 0.0
        naam_el = getattr(el, 'name', '-') if el is not None else '-'
        profiel = naam_el.split('(')[0].strip() if el is not None else '-'
        lengte = abs((top or 0.0) - bottom) if el is not None else None

        rijen: list[tuple[str, str, str, str]] = [
            ('Grondkering', '', '', ''),
            ('Profiel', profiel, '[-]', ''),
            ('Staalkwaliteit', getattr(el, 'steel_quality', '-') if el else '-', '[-]', ''),
            (
                'Opneembaar moment',
                fmt_number(getattr(el, 'opneembaar_moment_knm', None)) if el else '-',
                '[kNm/m]',
                '',
            ),
            ('Niveau b.k.', fmt_number(top or 0.0, 2) if el else '-', '[m NAP]', ''),
            ('Niveau o.k.', fmt_number(bottom, 2) if el else '-', '[m NAP]', ''),
            ('Lengte', fmt_number(lengte) if el else '-', '[m]', ''),
            ('Resultaten', '', '', 'Verificatiestap'),
            (
                'Moment Msd UGT',
                fmt_number(msd[0]),
                '[kNm/m]',
                self._formatteer_stap(msd[1], _step_short_label),
            ),
            (
                'Dwarskracht Dsd UGT',
                fmt_number(dsd[0]),
                '[kN/m]',
                self._formatteer_stap(dsd[1], _step_short_label),
            ),
            (
                'Verplaatsing urep BGT',
                fmt_number(vervorming[0]),
                '[mm]',
                self._formatteer_stap(vervorming[1], _step_short_label),
            ),
            ('Gemobiliseerd Moment', fmt_number(max_mob_mom), '[%]', ''),
            ('Gemobiliseerd Grond', fmt_number(max_mob_grond), '[%]', ''),
        ]

        # Ondersteuningen — identieke logica als de resultaattabel in de app
        anker_max: dict[str, tuple[float, str]] = {}
        for item in anchor_strut_resume:
            stap = _VTYPE_LABELS.get(item.verification_type, '')
            if item.name not in anker_max or abs(item.force) > abs(anker_max[item.name][0]):
                anker_max[item.name] = (item.force, stap)
        for item in supports_resume:
            stap = _VTYPE_LABELS.get(item.verification_type, '')
            if item.name not in anker_max or abs(item.force) > abs(anker_max[item.name][0]):
                anker_max[item.name] = (item.force, stap)

        steun_moment_max: dict[str, tuple[float, str]] = {}
        steun_heeft_kracht: set[str] = set()
        steun_heeft_moment: set[str] = set()
        for item in supports_resume:
            stap = _VTYPE_LABELS.get(item.verification_type, '')
            if abs(item.force) > 1e-6:
                steun_heeft_kracht.add(item.name)
            if abs(item.moment) > 1e-6:
                steun_heeft_moment.add(item.name)
            if item.name not in steun_moment_max or abs(item.moment) > abs(steun_moment_max[item.name][0]):
                steun_moment_max[item.name] = (item.moment, stap)

        niveau_per_naam: dict[str, float] = {}
        for rs in summaries:
            for naam_n, _kracht, niveau in rs.ondersteuningen:
                if naam_n not in niveau_per_naam:
                    niveau_per_naam[naam_n] = niveau
        for sp in spring_supports:
            if sp.name not in niveau_per_naam:
                niveau_per_naam[sp.name] = sp.level
        for rp in rigid_supports:
            if rp.name not in niveau_per_naam:
                niveau_per_naam[rp.name] = rp.level

        if anker_max:
            rijen.append(('Resultaten ondersteuningen', '', '', 'Verificatiestap'))
            for naam in sorted(anker_max):
                max_kracht, stap_kracht = anker_max[naam]
                stap_str = f'stap {stap_kracht}' if stap_kracht else ''
                if naam in niveau_per_naam:
                    rijen.append((f'Niveau {naam}', fmt_number(niveau_per_naam[naam], 2), '[m NAP]', ''))
                is_steun = naam in steun_moment_max
                toon_kracht = not is_steun or naam in steun_heeft_kracht
                toon_moment = is_steun and naam in steun_heeft_moment
                if toon_kracht:
                    rijen.append((naam, fmt_number(abs(max_kracht)), '[kN/m]', stap_str))
                if toon_moment:
                    max_moment, stap_moment = steun_moment_max[naam]
                    stap_str_m = f'stap {stap_moment}' if stap_moment else ''
                    rijen.append((naam, fmt_number(abs(max_moment)), '[kNm/m]', stap_str_m))

        return rijen

    def _stel_verticale_celmerge_in(self, cel_top, cel_bot) -> None:
        """Voeg verticale w:vMerge samen voor twee opeenvolgende cellen."""
        tc_pr_top = cel_top._tc.get_or_add_tcPr()
        v_merge_top = OxmlElement('w:vMerge')
        v_merge_top.set(qn('w:val'), 'restart')
        tc_pr_top.append(v_merge_top)
        v_align = OxmlElement('w:vAlign')
        v_align.set(qn('w:val'), 'top')
        tc_pr_top.append(v_align)

        tc_pr_bot = cel_bot._tc.get_or_add_tcPr()
        v_merge_bot = OxmlElement('w:vMerge')
        tc_pr_bot.append(v_merge_bot)
        for para in cel_bot.paragraphs:
            for run in para.runs:
                run.text = ''

    def _formatteer_stap(self, stap_key: str | None, label_fn) -> str:
        """Formatteer een verificatiestap voor de resultaat-specificatietabel."""
        return f'stap {label_fn(stap_key)}' if stap_key else ''

    def _maatgevende_resultaatwaarden(
        self,
        project: object,
        is_ugt_step_key,
        is_bgt_step_key,
    ) -> tuple[float | None, float | None, float | None]:
        """Bereken projectbrede Msd, Dsd en BGT-vervorming voor Word."""
        msd, dsd, vervorming = self._maatgevende_resultaatwaarden_en_stappen(
            project,
            is_ugt_step_key,
            is_bgt_step_key,
        )
        return msd[0], dsd[0], vervorming[0]

    def _maatgevende_resultaatwaarden_en_stappen(
        self,
        project: object,
        is_ugt_step_key,
        is_bgt_step_key,
    ) -> tuple[
        tuple[float | None, str | None],
        tuple[float | None, str | None],
        tuple[float | None, str | None],
    ]:
        """Bereken projectbrede resultaatwaarden inclusief stap-sleutel."""
        result_steps = getattr(project, 'result_steps', None) or {}
        msd: float | None = None
        dsd: float | None = None
        msd_stap: str | None = None
        dsd_stap: str | None = None
        vervorming: float | None = None
        vervorming_stap: str | None = None

        for stap_key, step in result_steps.items():
            if not is_ugt_step_key(stap_key):
                continue
            for rs in step.stages.values():
                for pt in rs.points:
                    moment = abs(pt.moment)
                    shear = abs(pt.shear)
                    if msd is None or moment > msd:
                        msd = moment
                        msd_stap = stap_key
                    if dsd is None or shear > dsd:
                        dsd = shear
                        dsd_stap = stap_key

        for stap_key, step in result_steps.items():
            if not is_bgt_step_key(stap_key):
                continue
            for rs in step.stages.values():
                for pt in rs.points:
                    disp = abs(pt.disp)
                    if vervorming is None or disp > vervorming:
                        vervorming = disp
                        vervorming_stap = stap_key
        return (msd, msd_stap), (dsd, dsd_stap), (vervorming, vervorming_stap)

    def _pas_resultaat_specificaties_opmaak_toe(self, tbl: Table) -> None:
        """Pas theme-opmaak toe op de resultaat-specificatietabel."""
        from ui import table_styles

        font = _eerste_fontfamilie(table_styles.TABLE_FONT)
        label_kleur = _hex_zonder_hash(table_styles.TABLE_LABEL_COLOR)
        waarde_kleur = _hex_zonder_hash(table_styles.TABLE_VALUE_COLOR)
        extra_kleur = _hex_zonder_hash(table_styles.TABLE_EXTRA_COLOR)
        header_bg = _hex_zonder_hash(table_styles.TABLE_HEADER_BG)
        header_fg = _hex_zonder_hash(table_styles.TABLE_HEADER_FG)
        row_odd_bg = _hex_zonder_hash(table_styles.TABLE_ROW_ODD_BG)
        row_even_bg = _hex_zonder_hash(table_styles.TABLE_ROW_EVEN_BG)

        data_index = 0
        for row in tbl.rows:
            eerste_tekst = row.cells[0].text.strip() if row.cells else ''
            is_koprij = (
                eerste_tekst in {'Grondkering', 'Damwand'}
                or eerste_tekst.startswith('Resultaten')
            )
            if is_koprij:
                for cell in row.cells:
                    self._stel_cel_vulling(cell, header_bg)
                    self._pas_cel_font_toe(
                        cell, font, header_fg, bold=True,
                        size_pt=table_styles.WORD_TABLE_HEADER_SIZE,
                    )
                continue

            fill = row_odd_bg if data_index % 2 == 0 else row_even_bg
            kleuren = [label_kleur, waarde_kleur, extra_kleur, extra_kleur]
            for col_idx, cell in enumerate(row.cells[:4]):
                self._stel_cel_vulling(cell, fill)
                self._pas_cel_font_toe(
                    cell, font, kleuren[min(col_idx, len(kleuren) - 1)], bold=False,
                    size_pt=table_styles.WORD_TABLE_TEXT_SIZE,
                )
            data_index += 1

    def _schrijf_extremen_overzicht_sectie(
        self,
        doc: Document,
        sec: ReportSection,
        project: object | None,
    ) -> None:
        """Schrijf de grafische maatgevende resultaten volgens het moederbestand."""
        doc.add_paragraph(RESULTATEN_GRAFIEK_INTRO_TEKST)
        doc.add_paragraph()
        for groep in sec.image_groups:
            self._schrijf_figuurgroep(doc, groep, project)
        doc.add_paragraph()
        if project is not None:
            doc.add_paragraph(self._resultaat_conclusie_tekst(project))

    def _resultaat_conclusie_tekst(self, project: object) -> str:
        """Bouw de conclusie onder de maatgevende-resultatenfiguur."""
        from reporting.builders.result_description_builder import (
            is_bgt_step_key,
            is_ugt_step_key,
        )
        from utils.formatting import fmt_number

        msd, _dsd, vervorming = self._maatgevende_resultaatwaarden(
            project,
            is_ugt_step_key,
            is_bgt_step_key,
        )
        sheet_piling = getattr(project, 'sheet_piling', None) or []
        opneembaar = (
            getattr(sheet_piling[0], 'opneembaar_moment_knm', None)
            if sheet_piling else None
        )
        moment_zin = (
            'de berekende momenten opneembaar zijn'
            if msd is None or opneembaar is None or abs(msd) <= opneembaar
            else 'de opneembaarheid van de berekende momenten nader beoordeeld moet worden'
        )
        return (
            'Op basis van bovenstaande resultaten kan worden geconcludeerd dat '
            f'{moment_zin}. De berekende topverplaatsing bij de doorsnede is '
            f'{fmt_number(vervorming)}mm. Op basis hiervan wordt voldaan aan de '
            'verplaatsingseis.'
        )

    def _schrijf_fase_sectie(
        self,
        doc: Document,
        sec: FaseInvoerSectie,
        project: object | None,
    ) -> None:
        """Schrijf een fase-sectie als tabel met invoer en doorsnede.

        Parameters
        ----------
        doc:
            Doeldocument.
        sec:
            Fase-sectie met ``fase_card``.
        project:
            Project voor het renderen van de doorsnede; mag ``None`` zijn.
        """
        kaart = sec.fase_card
        if kaart is None:
            return

        # Render afbeelding en bereken tabelinhoud vóór pagineringsbeslissing
        png_bytes: bytes | None = None
        if project is not None:
            img_req = ReportImageRequest(
                id=f'fase_{kaart.fase_num}_doorsnede',
                caption='',
                figure_key='section',
                stage_index=kaart.fase_num - 1,
                step_key=None,
            )
            png_bytes = render_figuur(img_req, project)

        n_data = sum(1 + len(rij.extra_lines) for rij in kaart.rows)
        afbeeldingshoogte_cm = _png_hoogte_cm(png_bytes, 6.0) if png_bytes else 0.0
        tekst_cm = sum(
            _FASE_RIJHOOGTE_CM * (1 + len(rij.extra_lines))
            for rij in kaart.rows
        ) if kaart.rows else _FASE_RIJHOOGTE_CM
        padding_cm = max(afbeeldingshoogte_cm * 1.05 - tekst_cm, 0)

        if self._fase_tabel_teller > 0:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)
            spacer.paragraph_format.keep_with_next = True
        n_padding = 1 if padding_cm > 0 else 0
        n_header = 2
        n_rows = n_header + max(n_data, 1) + n_padding

        tbl = doc.add_table(rows=n_rows, cols=4)
        tbl.autofit = False
        try:
            tbl.style = 'Table Grid'
        except KeyError:
            pass

        kolom_breedtes = [1701, 1134, 2835, 3572]
        self._stel_tabel_grid_in(tbl, [3.0, 2.0, 5.0, 6.3])
        for row in tbl.rows:
            for col_idx, cell in enumerate(row.cells[:4]):
                self._stel_cel_breedte(cell, kolom_breedtes[col_idx])

        cel_naam = tbl.rows[0].cells[0].merge(tbl.rows[0].cells[2])
        self._stel_cel_breedte(cel_naam, sum(kolom_breedtes[:3]))
        cel_naam.text = kaart.stage_name
        tbl.rows[0].cells[3].text = 'Afbeelding'

        tbl.rows[1].cells[0].text = 'Parameter'
        tbl.rows[1].cells[1].text = 'Niveau'
        tbl.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tbl.rows[1].cells[2].text = 'Toelichting'
        tbl.rows[1].cells[3].text = ''

        grid_row = n_header
        for rij in kaart.rows:
            n_sub = 1 + len(rij.extra_lines)
            if n_sub > 1:
                label_cel = tbl.rows[grid_row].cells[0].merge(
                    tbl.rows[grid_row + n_sub - 1].cells[0]
                )
                niveau_cel = tbl.rows[grid_row].cells[1].merge(
                    tbl.rows[grid_row + n_sub - 1].cells[1]
                )
                self._stel_cel_breedte(label_cel, kolom_breedtes[0])
                self._stel_cel_breedte(niveau_cel, kolom_breedtes[1])
                for k in range(n_sub):
                    self._stel_rijhoogte_exact_twips(tbl.rows[grid_row + k], _DAMWAND_RIJHOOGTE_TWIPS)
            else:
                self._stel_rijhoogte_exact_twips(tbl.rows[grid_row], _DAMWAND_RIJHOOGTE_TWIPS)

            tbl.rows[grid_row].cells[0].text = rij.label
            tbl.rows[grid_row].cells[1].text = rij.value
            tbl.rows[grid_row].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            tbl.rows[grid_row].cells[2].text = rij.extra
            for k, extra_tekst in enumerate(rij.extra_lines):
                tbl.rows[grid_row + k + 1].cells[2].text = extra_tekst
            grid_row += n_sub

        if n_data == 0:
            self._stel_rijhoogte_exact_twips(tbl.rows[n_header], _DAMWAND_RIJHOOGTE_TWIPS)

        if n_padding:
            pad_start = n_header + max(n_data, 1)
            pad_cel = tbl.rows[pad_start].cells[0].merge(tbl.rows[n_rows - 1].cells[2])
            self._stel_cel_breedte(pad_cel, sum(kolom_breedtes[:3]))
            self._stel_rijhoogte_exact_twips(tbl.rows[pad_start], round(padding_cm * 567))

        self._pas_fase_tabel_opmaak_toe(tbl)
        for row in tbl.rows:
            self._stel_cant_split(row)

        img_cel = tbl.rows[n_header].cells[3]
        for rij_idx in range(n_header + 1, n_rows):
            img_cel = img_cel.merge(tbl.rows[rij_idx].cells[3])
        self._stel_cel_breedte(img_cel, kolom_breedtes[3])

        para = img_cel.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        if png_bytes:
            run = para.add_run()
            run.add_picture(io.BytesIO(png_bytes), width=Cm(6))

        tc_pr = img_cel._tc.get_or_add_tcPr()
        v_align = OxmlElement('w:vAlign')
        v_align.set(qn('w:val'), 'center')
        tc_pr.append(v_align)
        self._houd_tabel_bij_elkaar(tbl)
        self._fase_tabel_teller = getattr(self, '_fase_tabel_teller', 0) + 1

    def _pas_fase_tabel_opmaak_toe(self, tbl: Table) -> None:
        """Pas thema-font en headerkleuren toe op de fase-invoertabel."""
        from ui import table_styles

        font = _eerste_fontfamilie(table_styles.TABLE_FONT)
        for row in tbl.rows:
            for cell in row.cells:
                self._pas_cel_font_toe(
                    cell, font, '000000', bold=False,
                    size_pt=table_styles.WORD_TABLE_TEXT_SIZE,
                )

        header_bg = _hex_zonder_hash(table_styles.TABLE_HEADER_BG)
        header_fg = _hex_zonder_hash(table_styles.TABLE_HEADER_FG)
        subheader_bg = _hex_zonder_hash(table_styles.TABLE_HEADER_SUB_BG)
        subheader_fg = _hex_zonder_hash(table_styles.TABLE_HEADER_SUB_FG)
        for cell in tbl.rows[0].cells:
            self._stel_cel_vulling(cell, header_bg)
            self._pas_cel_font_toe(
                cell, font, header_fg, bold=True,
                size_pt=table_styles.WORD_TABLE_HEADER_SIZE,
            )
        for cell in tbl.rows[1].cells:
            self._stel_cel_vulling(cell, subheader_bg)
            self._pas_cel_font_toe(
                cell, font, subheader_fg, bold=True,
                size_pt=table_styles.WORD_TABLE_HEADER_SIZE,
            )

    def _houd_tabel_bij_elkaar(self, tbl: Table) -> None:
        """Laat Word een fasetabel zoveel mogelijk als één blok pagineren.

        Itereert via raw XML zodat ook vMerge-stub cellen worden meegenomen;
        python-docx ``row.cells`` geeft bij verticale celfusies alleen de
        restart-cel terug en slaat de stubs over.
        """
        laatste_rij_idx = len(tbl.rows) - 1
        for row_idx, row in enumerate(tbl.rows):
            kwn = row_idx < laatste_rij_idx
            for tc in row._tr.findall(qn('w:tc')):
                for p in tc.findall(qn('w:p')):
                    p_pr = p.find(qn('w:pPr'))
                    if p_pr is None:
                        p_pr = OxmlElement('w:pPr')
                        p.insert(0, p_pr)
                    if p_pr.find(qn('w:keepLines')) is None:
                        p_pr.append(OxmlElement('w:keepLines'))
                    keep_next = p_pr.find(qn('w:keepNext'))
                    if kwn and keep_next is None:
                        p_pr.append(OxmlElement('w:keepNext'))
                    elif not kwn and keep_next is not None:
                        p_pr.remove(keep_next)

    def _pas_cel_font_toe(
        self,
        cell,
        font_name: str,
        color_hex: str,
        *,
        bold: bool,
        size_pt: int,
    ) -> None:
        """Zet font, kleur en gewicht voor alle runs in een Word-cel."""
        kleur = RGBColor.from_string(color_hex)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(size_pt)
                run.font.bold = bold
                run.font.color.rgb = kleur
                r_pr = run._r.get_or_add_rPr()
                r_fonts = r_pr.rFonts
                if r_fonts is None:
                    r_fonts = OxmlElement('w:rFonts')
                    r_pr.append(r_fonts)
                r_fonts.set(qn('w:ascii'), font_name)
                r_fonts.set(qn('w:hAnsi'), font_name)

    def _stel_cel_vulling(self, cell, fill_hex: str) -> None:
        """Zet celvulling op een hexkleur zonder randstijlen te wijzigen."""
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            tc_pr.append(shd)
        shd.set(qn('w:fill'), fill_hex)

    def _verwijder_cel_randen(self, cell) -> None:
        """Zet alle celranden op 'none' (onzichtbaar)."""
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)
        for kant in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            rand = OxmlElement(f'w:{kant}')
            rand.set(qn('w:val'), 'none')
            rand.set(qn('w:sz'), '0')
            rand.set(qn('w:space'), '0')
            rand.set(qn('w:color'), 'auto')
            tc_borders.append(rand)

    def _stel_cel_breedte(self, cell, breedte_dxa: int) -> None:
        """Zet celbreedte in twips/dxa zoals Word die in ``tcW`` verwacht."""
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn('w:tcW'))
        if tc_w is None:
            tc_w = OxmlElement('w:tcW')
            tc_pr.append(tc_w)
        tc_w.set(qn('w:w'), str(breedte_dxa))
        tc_w.set(qn('w:type'), 'dxa')

    def _stel_cel_verticale_uitlijning(self, cell, waarde: str) -> None:
        """Zet verticale celuitlijning in Word."""
        tc_pr = cell._tc.get_or_add_tcPr()
        v_align = tc_pr.find(qn('w:vAlign'))
        if v_align is None:
            v_align = OxmlElement('w:vAlign')
            tc_pr.append(v_align)
        v_align.set(qn('w:val'), waarde)

    def _stel_cant_split(self, row) -> None:
        """Voorkom dat Word deze rij over een pagina-einde breekt."""
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn('w:cantSplit')) is None:
            cant_split = OxmlElement('w:cantSplit')
            tr_pr.append(cant_split)

    def _stel_rijhoogte_twips(self, row, hoogte_twips: int) -> None:
        """Zet een compacte rijhoogte in Word-twips zonder exact-height regel."""
        tr_pr = row._tr.get_or_add_trPr()
        tr_height = tr_pr.find(qn('w:trHeight'))
        if tr_height is None:
            tr_height = OxmlElement('w:trHeight')
            tr_pr.append(tr_height)
        tr_height.set(qn('w:val'), str(hoogte_twips))

    def _stel_rijhoogte_exact_twips(self, row, hoogte_twips: int) -> None:
        """Zet een vaste rijhoogte in Word-twips."""
        tr_pr = row._tr.get_or_add_trPr()
        tr_height = tr_pr.find(qn('w:trHeight'))
        if tr_height is None:
            tr_height = OxmlElement('w:trHeight')
            tr_pr.append(tr_height)
        tr_height.set(qn('w:val'), str(hoogte_twips))
        tr_height.set(qn('w:hRule'), 'exact')

    def _stel_tabel_grid_in(self, tbl: Table, breedtes_cm: list[float]) -> None:
        """Stel vaste Word-kolombreedtes in via ``tblGrid``."""
        tbl_grid = tbl._tbl.find(qn('w:tblGrid'))
        if tbl_grid is None:
            tbl_grid = OxmlElement('w:tblGrid')
            tbl._tbl.insert(0, tbl_grid)
        else:
            for col in list(tbl_grid):
                tbl_grid.remove(col)

        for breedte_cm in breedtes_cm:
            grid_col = OxmlElement('w:gridCol')
            grid_col.set(qn('w:w'), str(round(breedte_cm * 567)))
            tbl_grid.append(grid_col)

    def _schrijf_tabel(self, doc: Document, tabel) -> None:
        if not tabel.columns:
            return
        if tabel.title:
            doc.add_paragraph(tabel.title)
        heeft_groepen = bool(tabel.column_groups)
        extra_rij = 1 if heeft_groepen else 0
        n_cols = len(tabel.columns)
        t = doc.add_table(rows=1 + extra_rij + len(tabel.rows), cols=n_cols)
        try:
            t.style = 'Table Grid'
        except KeyError:
            pass
        vaste_breedtes_cm = (
            _GRONDSOORTEN_V2_FASE_KOLOM_BREEDTES_CM
            if str(tabel.id).startswith('grondsoorten_v2_fase_')
            and n_cols == len(_GRONDSOORTEN_V2_FASE_KOLOM_BREEDTES_CM)
            else []
        )
        t.autofit = not bool(vaste_breedtes_cm)
        vaste_breedtes_dxa = [round(cm * 567) for cm in vaste_breedtes_cm]
        if vaste_breedtes_cm:
            self._stel_tabel_grid_in(t, vaste_breedtes_cm)
            for row in t.rows:
                for col_idx, cell in enumerate(row.cells[:n_cols]):
                    self._stel_cel_breedte(cell, vaste_breedtes_dxa[col_idx])

        if heeft_groepen:
            col_offset = 0
            for groep_label, colspan in tabel.column_groups:
                start = t.rows[0].cells[col_offset]
                if colspan > 1:
                    eind = t.rows[0].cells[col_offset + colspan - 1]
                    start = start.merge(eind)
                    if vaste_breedtes_dxa:
                        self._stel_cel_breedte(
                            start,
                            sum(vaste_breedtes_dxa[col_offset:col_offset + colspan]),
                        )
                start.text = groep_label
                col_offset += colspan

        kop_rij = extra_rij
        for col, header in enumerate(tabel.columns):
            cell = t.rows[kop_rij].cells[col]
            if '\n' in header:
                parts = header.split('\n')
                cell.text = parts[0]
                for part in parts[1:]:
                    cell.add_paragraph(part)
            else:
                cell.text = header

        for row_i, data_rij in enumerate(tabel.rows):
            rij = t.rows[kop_rij + 1 + row_i]
            for col, cel in enumerate(data_rij):
                if col < len(rij.cells):
                    rij.cells[col].text = str(cel)

        if vaste_breedtes_dxa:
            self._voeg_grondsoorten_v2_ongewijzigd_merges_toe(
                t,
                tabel.rows,
                kop_rij + 1,
                vaste_breedtes_dxa,
            )

        self._pas_report_tabel_opmaak_toe(t, heeft_groepen)
        links_uitlijnen = self._grondsoorten_v2_laagkolommen(tabel.id)
        if links_uitlijnen:
            self._lijn_report_datarij_kolommen_links(
                t,
                kop_rij + 1,
                links_uitlijnen,
            )

    def _voeg_grondsoorten_v2_ongewijzigd_merges_toe(
        self,
        tbl: Table,
        data_rijen: list[list[str]],
        data_start: int,
        kolom_breedtes_dxa: list[int],
    ) -> None:
        """Voeg ongewijzigde v2-grondopbouwzijden samen zoals in de app."""
        if not data_rijen:
            return

        melding = 'Grondopbouw ongewijzigd t.o.v. vorige fase'
        for col_start in (0, 3):
            eerste_rij = data_rijen[0]
            if col_start + 2 >= len(eerste_rij) or eerste_rij[col_start] != melding:
                continue
            zijde_is_blok = eerste_rij[col_start + 1:col_start + 3] == ['', '']
            zijde_is_blok = zijde_is_blok and all(
                rij[col_start:col_start + 3] == ['', '', '']
                for rij in data_rijen[1:]
                if col_start + 2 < len(rij)
            )
            if not zijde_is_blok:
                continue

            start = tbl.rows[data_start].cells[col_start]
            eind = tbl.rows[data_start + len(data_rijen) - 1].cells[col_start + 2]
            samengevoegd = start.merge(eind)
            self._stel_cel_breedte(
                samengevoegd,
                sum(kolom_breedtes_dxa[col_start:col_start + 3]),
            )
            samengevoegd.text = melding
            self._stel_cel_verticale_uitlijning(samengevoegd, 'center')

    def _grondsoorten_v2_laagkolommen(self, tabel_id: str) -> list[int]:
        """Geef links uit te lijnen laagnaamkolommen voor v2-tabellen."""
        if tabel_id == 'grondsoorten_v2_overzicht_tabel':
            return [0]
        if tabel_id.startswith('grondsoorten_v2_fase_'):
            return [0, 3]
        return []

    def _lijn_report_datarij_kolommen_links(
        self,
        tbl: Table,
        data_start: int,
        kolommen: list[int],
    ) -> None:
        """Lijn specifieke datarijkolommen in een rapporttabel links uit."""
        for row in tbl.rows[data_start:]:
            for col in kolommen:
                if col < len(row.cells):
                    row.cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _pas_report_tabel_opmaak_toe(self, tbl: Table, heeft_groepen: bool = False) -> None:
        """Pas themakleuren toe op een generieke rapport-tabel."""
        from ui import table_styles

        font = _eerste_fontfamilie(table_styles.TABLE_FONT)
        header_bg = _hex_zonder_hash(table_styles.TABLE_HEADER_BG)
        header_fg = _hex_zonder_hash(table_styles.TABLE_HEADER_FG)
        subhdr_bg = _hex_zonder_hash(table_styles.TABLE_HEADER_SUB_BG)
        subhdr_fg = _hex_zonder_hash(table_styles.TABLE_HEADER_SUB_FG)
        value_kleur = _hex_zonder_hash(table_styles.TABLE_VALUE_COLOR)
        row_odd_bg = _hex_zonder_hash(table_styles.TABLE_ROW_ODD_BG)
        row_even_bg = _hex_zonder_hash(table_styles.TABLE_ROW_EVEN_BG)
        kop_rij = 1 if heeft_groepen else 0

        if heeft_groepen:
            for cell in tbl.rows[0].cells:
                self._stel_cel_vulling(cell, subhdr_bg)
                self._pas_cel_font_toe(
                    cell, font, subhdr_fg, bold=True,
                    size_pt=table_styles.WORD_TABLE_HEADER_SIZE,
                )

        for cell in tbl.rows[kop_rij].cells:
            self._stel_cel_vulling(cell, header_bg)
            self._pas_cel_font_toe(
                cell, font, header_fg, bold=True,
                size_pt=table_styles.WORD_TABLE_HEADER_SIZE,
            )
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row_i, row in enumerate(tbl.rows[kop_rij + 1:]):
            fill = row_odd_bg if row_i % 2 == 0 else row_even_bg
            for cell in row.cells:
                self._stel_cel_vulling(cell, fill)
                self._pas_cel_font_toe(
                    cell, font, value_kleur, bold=False,
                    size_pt=table_styles.WORD_TABLE_TEXT_SIZE,
                )
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _schrijf_figuurgroep(
        self,
        doc: Document,
        groep: ReportImageGroup,
        project: object | None,
    ) -> None:
        """Schrijf een ReportImageGroup als 3-rijen Word-tabel (kop, figuur, voet)."""
        if not groep.images:
            return
        n_cols = len(groep.images)
        breedte_per_col = 16.0 / n_cols
        breedte_dxa = round(breedte_per_col * 567)

        tbl = doc.add_table(rows=3, cols=n_cols)
        try:
            tbl.style = 'Table Grid'
        except KeyError:
            pass
        self._stel_tabel_grid_in(tbl, [breedte_per_col] * n_cols)

        self._stel_rijhoogte_exact_twips(tbl.rows[0], _DAMWAND_RIJHOOGTE_TWIPS)
        for col, tekst in enumerate(groep.headers):
            cell = tbl.rows[0].cells[col]
            self._stel_cel_breedte(cell, breedte_dxa)
            cell.text = tekst
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for col, img_req in enumerate(groep.images):
            cell = tbl.rows[1].cells[col]
            self._stel_cel_breedte(cell, breedte_dxa)
            if img_req is not None and project is not None:
                png = self._render_figuur(img_req, project)
                if png:
                    para = cell.paragraphs[0]
                    run = para.add_run()
                    run.add_picture(io.BytesIO(png), width=Cm(breedte_per_col - 0.4))

        self._stel_rijhoogte_exact_twips(tbl.rows[2], _DAMWAND_RIJHOOGTE_TWIPS)
        for col, tekst in enumerate(groep.footers):
            cell = tbl.rows[2].cells[col]
            self._stel_cel_breedte(cell, breedte_dxa)
            cell.text = tekst
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._pas_figuurgroep_opmaak_toe(tbl)

    def _pas_figuurgroep_opmaak_toe(self, tbl: Table) -> None:
        """Pas themakleuren toe op een figuurgroep-tabel."""
        from ui import table_styles

        font = _eerste_fontfamilie(table_styles.TABLE_FONT)
        header_bg = _hex_zonder_hash(table_styles.TABLE_HEADER_BG)
        header_fg = _hex_zonder_hash(table_styles.TABLE_HEADER_FG)
        row_odd_bg = _hex_zonder_hash(table_styles.TABLE_ROW_ODD_BG)
        extra_kleur = _hex_zonder_hash(table_styles.TABLE_EXTRA_COLOR)

        for cell in tbl.rows[0].cells:
            self._stel_cel_vulling(cell, header_bg)
            self._pas_cel_font_toe(
                cell, font, header_fg, bold=True,
                size_pt=table_styles.WORD_TABLE_HEADER_SIZE,
            )
        for cell in tbl.rows[1].cells:
            self._stel_cel_vulling(cell, row_odd_bg)
        for cell in tbl.rows[2].cells:
            self._stel_cel_vulling(cell, row_odd_bg)
            self._pas_cel_font_toe(
                cell, font, extra_kleur, bold=False,
                size_pt=table_styles.WORD_TABLE_TEXT_SIZE,
            )

    # ------------------------------------------------------------------
    # Figuren
    # ------------------------------------------------------------------

    def _schrijf_figuur(self, doc: Document, img_req, project) -> None:
        """Render een figuur headless en voeg in als inline afbeelding."""
        if project is None:
            doc.add_paragraph(f'[Figuur: {img_req.caption}]')
            return
        png_bytes = self._render_figuur(img_req, project)
        if png_bytes:
            doc.add_paragraph(img_req.caption)
            doc.add_picture(io.BytesIO(png_bytes), width=Cm(16))
        else:
            doc.add_paragraph(f'[Figuur niet beschikbaar: {img_req.caption}]')

    def _render_figuur(self, img_req, project) -> bytes | None:
        """Render figuur naar PNG-bytes met headless matplotlib.

        Ondersteunde figure_key waarden:
        - 'section'      : dwarsdoorsnede via SectionRenderer
        - 'moment_shear' : moment + dwarskracht via OutputRenderer
        - 'displacement' : vervorming via OutputRenderer
        """
        return render_figuur(img_req, project)
