"""Tests voor WordExporter."""

from __future__ import annotations

import base64
import os
import tempfile

from docx import Document

from exporters import word_exporter
from exporters.word_exporter import WordExporter
from reporting.models import (
    FaseInvoerSectie,
    ReportField,
    ReportImageGroup,
    ReportImageRequest,
    ReportItem,
    ReportMetadata,
    ReportPackage,
    ReportSection,
)
from reporting.builders.input_description_builder import FaseCard, FaseRow


_PNG_1X1 = base64.b64decode(
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwAD'
    'hgGAWjR9awAAAABJRU5ErkJggg=='
)


def test_word_exporter_schrijft_image_als_inline_shape(monkeypatch) -> None:
    """Een sectie met image-request geeft een inline shape in de docx."""
    monkeypatch.setattr(word_exporter, 'render_figuur', lambda _img, _project: _PNG_1X1)
    sec = ReportSection(
        id='fase_1',
        title='Fase 1',
        images=[
            ReportImageRequest(
                id='img_fase_1',
                caption='Doorsnede fase 1',
                figure_key='section',
                stage_index=0,
                step_key=None,
            )
        ],
    )
    pkg = ReportPackage(
        metadata=ReportMetadata(project_name='T'),
        input_sections=[sec],
        selected_items=[
            ReportItem(
                id='damwand_fase_1',
                kind='invoer',
                caption='Fase 1',
                source_ref='fase_1',
            )
        ],
    )
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as handle:
        out = handle.name
    fout = WordExporter().export(pkg, None, out, project=object())
    assert fout is None
    doc = Document(out)
    os.unlink(out)
    assert len(doc.inline_shapes) == 1


def test_word_exporter_schrijft_image_group_als_tabel(monkeypatch) -> None:
    """Een 3x3 figuurgroep wordt als Word-tabel met drie afbeeldingen geschreven."""
    monkeypatch.setattr(word_exporter, 'render_figuur', lambda _img, _project: _PNG_1X1)
    img = ReportImageRequest(
        id='m',
        caption='',
        figure_key='moment_curve',
        stage_index=0,
        step_key='CUR 166 6.4',
    )
    sec = ReportSection(
        id='extremen_overzicht',
        title='Maatgevende resultaten',
        image_groups=[
            ReportImageGroup(
                id='extremen_3x3',
                title='',
                headers=['Msd = 210 kNm/m', 'Dsd = 95 kN/m', 'Urep BGT = 12 mm'],
                images=[img, img, img],
                footers=['Fase 1 - Start', 'Fase 2 - Eind', 'Fase 2 - Eind'],
            )
        ],
    )
    pkg = ReportPackage(
        metadata=ReportMetadata(project_name='T'),
        result_sections=[sec],
        selected_items=[
            ReportItem(
                id='result_extremen_overzicht',
                kind='resultaat',
                caption='Maatgevende resultaten',
                source_ref='extremen_overzicht',
            )
        ],
    )
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as handle:
        out = handle.name
    fout = WordExporter().export(pkg, None, out, project=object())
    assert fout is None
    doc = Document(out)
    os.unlink(out)
    assert len(doc.tables) == 2
    assert len(doc.inline_shapes) == 3
    assert doc.tables[-1].rows[0].cells[0].text == 'Msd = 210 kNm/m'


def test_word_exporter_gebruikt_fase_invoer_tabel_layout() -> None:
    """Normale Word-export gebruikt dezelfde fase-layout als hoofdstukexport."""
    kaart = FaseCard(fase_num=2, stage_name='Fase 2: Belasting')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    kaart.rows.append(FaseRow(
        'Bovenbelasting',
        'op maaiveld',
        '5,0 [kN/m²]',
        extra_lines=['3,0m breed', '0,0m vanaf damwand'],
    ))
    sec = FaseInvoerSectie(
        id='fase_2_invoer',
        title='Fase 2: Belasting',
        fase_card=kaart,
    )
    pkg = ReportPackage(
        metadata=ReportMetadata(project_name='T'),
        input_sections=[sec],
        selected_items=[
            ReportItem(
                id='damwand_fase_2_invoer',
                kind='invoer',
                caption='Fase 2: Belasting',
                source_ref='fase_2_invoer',
            )
        ],
    )
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as handle:
        out = handle.name
    fout = WordExporter().export(pkg, None, out, project=None)
    assert fout is None
    doc = Document(out)
    os.unlink(out)

    assert len(doc.tables) >= 2
    fase_tabel = doc.tables[-1]
    assert [c.text for c in fase_tabel.rows[1].cells][:3] == [
        'Parameter', 'Niveau', 'Toelichting',
    ]
    alle_tekst = '\n'.join(c.text for row in fase_tabel.rows for c in row.cells)
    assert 'Bovenbelasting' in alle_tekst
    assert '3,0m breed' in alle_tekst
    assert '0,0m vanaf damwand' in alle_tekst


def test_word_exporter_gebruikt_damwandgegevens_tabel_layout() -> None:
    """Normale Word-export schrijft damwandgegevens als voorbeeldtabel."""
    sec = ReportSection(id='damwand_gegevens', title='Damwandgegevens')
    sec.fields.extend([
        ReportField('profiel', 'Profiel', 'AZ 18-700'),
        ReportField('hoogte', 'Hoogte', '420,0', 'mm'),
    ])
    pkg = ReportPackage(
        metadata=ReportMetadata(project_name='T'),
        input_sections=[sec],
        selected_items=[
            ReportItem(
                id='damwand_damwand_gegevens',
                kind='invoer',
                caption='Damwandgegevens',
                source_ref='damwand_gegevens',
            )
        ],
    )
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as handle:
        out = handle.name
    fout = WordExporter().export(pkg, None, out, project=None)
    assert fout is None
    doc = Document(out)
    os.unlink(out)

    damwand_tabel = doc.tables[-1]
    assert [c.text for c in damwand_tabel.rows[0].cells] == [
        'Parameter', 'Waarde', 'Eenheid',
    ]
    assert [c.text for c in damwand_tabel.rows[2].cells] == [
        'Hoogte', '420,0', '[mm]',
    ]
