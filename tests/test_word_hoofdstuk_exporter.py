"""Tests voor WordHoofdstukExporter."""
from __future__ import annotations
import sys, os, tempfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from docx import Document
from reporting.models import ReportSection, ReportField, ReportTable, ReportImageRequest, ReportMetadata
from exporters.word_hoofdstuk_exporter import WordHoofdstukExporter


TEMPLATE = os.path.join(os.path.dirname(__file__), '..', 'templates', 'damwand_stijlen.docx')


def _export(secties, metadata=None) -> Document:
    """Exporteer naar een tijdelijk bestand en lees terug."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    exp = WordHoofdstukExporter()
    fout = exp.export(
        sections=secties,
        metadata=metadata or ReportMetadata(project_name='Testproject'),
        project=None,
        template_path=TEMPLATE,
        output_path=pad,
    )
    assert fout is None, f'Export fout: {fout}'
    doc = Document(pad)
    os.unlink(pad)
    return doc


def test_export_maakt_bestand_aan() -> None:
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    exp = WordHoofdstukExporter()
    fout = exp.export(
        sections=[],
        metadata=ReportMetadata(project_name='P'),
        project=None,
        template_path=TEMPLATE,
        output_path=pad,
    )
    assert fout is None
    assert os.path.exists(pad)
    os.unlink(pad)


def test_sectietitel_wordt_heading() -> None:
    sec = ReportSection(id='test', title='Mijn Sectie')
    doc = _export([sec])
    teksten = [p.text for p in doc.paragraphs]
    assert 'Mijn Sectie' in teksten


def test_veld_wordt_paragraaf() -> None:
    sec = ReportSection(id='test', title='Sectie')
    sec.fields.append(ReportField('k', 'Profiel', 'AZ 14-700', ''))
    doc = _export([sec])
    tekst = ' '.join(p.text for p in doc.paragraphs)
    assert 'Profiel' in tekst
    assert 'AZ 14-700' in tekst


def test_tabel_kolommen_aanwezig() -> None:
    sec = ReportSection(id='test', title='Sectie')
    sec.tables.append(ReportTable(
        id='t', title='',
        columns=['Fase', 'Moment'],
        rows=[['F1', '100']],
    ))
    doc = _export([sec])
    assert len(doc.tables) >= 1
    header_cellen = [c.text for c in doc.tables[0].rows[0].cells]
    assert 'Fase' in header_cellen
    assert 'Moment' in header_cellen


def test_export_zonder_template_gebruikt_lege_doc() -> None:
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    fout = WordHoofdstukExporter().export(
        sections=[ReportSection(id='x', title='Y')],
        metadata=ReportMetadata(),
        project=None,
        template_path=None,
        output_path=pad,
    )
    assert fout is None
    os.unlink(pad)


def test_png_hoogte_cm_berekent_correct() -> None:
    import struct
    import zlib

    from exporters.word_hoofdstuk_exporter import _png_hoogte_cm

    def _mini_png(w: int, h: int) -> bytes:
        header = b'\x89PNG\r\n\x1a\n'
        ihdr_data = struct.pack('>IIBBBBB', w, h, 8, 2, 0, 0, 0)
        crc = zlib.crc32(b'IHDR' + ihdr_data) & 0xFFFFFFFF
        ihdr = struct.pack('>I', 13) + b'IHDR' + ihdr_data + struct.pack('>I', crc)
        return header + ihdr

    png = _mini_png(100, 200)
    result = _png_hoogte_cm(png, 6.0)
    assert abs(result - 12.0) < 0.001


def test_fase_invoer_sectie_maakt_tabel_in_word() -> None:
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    kaart = FaseCard(fase_num=1, stage_name='Fase 1: Test')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    kaart.rows.append(FaseRow('Leganker', '-0,3 [m NAP]', '0 graden t.o.v. maaiveld'))
    sec = FaseInvoerSectie(id='fase_1', title='Fase 1: Test', fase_card=kaart)

    doc = _export([sec], metadata=ReportMetadata(project_name='Test'))
    assert len(doc.tables) >= 1


def test_damwandgegevens_sectie_wordt_voorbeeldtabel() -> None:
    from docx.oxml.ns import qn

    sec = ReportSection(id='damwand_gegevens', title='Damwandgegevens')
    sec.fields.extend([
        ReportField('profiel', 'Profiel', 'AZ 18-700'),
        ReportField('staalkwaliteit', 'Staalkwaliteit', 'S355'),
        ReportField('hoogte', 'Hoogte', '420,0', 'mm'),
        ReportField('ondersteuning_3', 'Anker A', '-2,1', 'm NAP'),
    ])

    doc = _export([sec], metadata=ReportMetadata(project_name='Test'))
    tbl = doc.tables[0]

    assert [col.get(qn('w:w')) for col in tbl._tbl.tblGrid.gridCol_lst] == [
        '2835', '1701', '1134',
    ]
    assert [cell.text for cell in tbl.rows[0].cells] == [
        'Parameter', 'Waarde', 'Eenheid',
    ]
    assert [cell.text for cell in tbl.rows[1].cells] == ['Profiel', 'AZ 18-700', '']
    assert [cell.text for cell in tbl.rows[3].cells] == ['Hoogte', '420,0', '[mm]']
    assert [cell.text for cell in tbl.rows[4].cells] == ['', '', '']
    assert [cell.text for cell in tbl.rows[5].cells] == [
        'Anker A', '-2,1', '[m NAP]',
    ]
    for cell in tbl.rows[0].cells:
        shd = cell._tc.tcPr.find(qn('w:shd'))
        assert shd is not None
        assert shd.get(qn('w:fill')) == '147ACF'
    assert tbl.rows[0]._tr.trPr.trHeight.get(qn('w:val')) == '255'
    assert tbl.rows[0]._tr.trPr.trHeight.get(qn('w:hRule')) == 'exact'
    assert tbl.rows[4]._tr.trPr.trHeight.get(qn('w:val')) == '40'


def test_fase_sectie_tabel_bevat_kolomhoofden() -> None:
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    kaart = FaseCard(fase_num=2, stage_name='Fase 2: Belasting')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    kaart.rows.append(FaseRow(
        'Bovenbelasting',
        'op maaiveld',
        '5,0 [kN/m²]',
        extra_lines=['3,0m breed', '0m vanaf damwand'],
    ))
    sec = FaseInvoerSectie(id='fase_2', title='Fase 2: Belasting', fase_card=kaart)

    doc = _export([sec], metadata=ReportMetadata())
    tbl = doc.tables[0]
    teksten_rij1 = [c.text for c in tbl.rows[1].cells]
    assert 'Parameter' in teksten_rij1
    assert 'Niveau' in teksten_rij1
    assert 'Toelichting' in teksten_rij1
    assert len(tbl.rows) >= 5


def test_fase_sectie_tabel_gebruikt_voorbeeld_breedtes_en_headerkleur() -> None:
    from docx.oxml.ns import qn
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    kaart = FaseCard(fase_num=2, stage_name='Fase 2: Belasting')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    kaart.rows.append(FaseRow(
        'Bovenbelasting',
        'op maaiveld',
        '5,0 [kN/m²]',
        extra_lines=['3,0m breed', '0,0m vanaf damwand'],
    ))
    sec = FaseInvoerSectie(id='fase_2', title='Fase 2: Belasting', fase_card=kaart)

    doc = _export([sec], metadata=ReportMetadata())
    tbl = doc.tables[0]
    grid = tbl._tbl.tblGrid
    assert [col.get(qn('w:w')) for col in grid.gridCol_lst] == [
        '1701', '1134', '2835', '3572',
    ]
    assert tbl.rows[0].cells[0]._tc.tcPr.tcW.w == 5670
    assert tbl.rows[0].cells[3]._tc.tcPr.tcW.w == 3572
    assert tbl.rows[1].cells[0]._tc.tcPr.tcW.w == 1701
    assert tbl.rows[1].cells[1]._tc.tcPr.tcW.w == 1134
    assert tbl.rows[1].cells[2]._tc.tcPr.tcW.w == 2835

    for row_idx in (0, 1):
        for cell in tbl.rows[row_idx].cells:
            shd = cell._tc.tcPr.find(qn('w:shd'))
            assert shd is not None
            assert shd.get(qn('w:fill')) == '147ACF'

    header_run = tbl.rows[1].cells[0].paragraphs[0].runs[0]
    assert header_run.font.size.pt == 8
    data_run = tbl.rows[2].cells[0].paragraphs[0].runs[0]
    assert data_run.font.size.pt == 7

    heights = [
        tbl.rows[i]._tr.trPr.trHeight.get(qn('w:val'))
        for i in range(len(tbl.rows))
        if tbl.rows[i]._tr.trPr is not None
        and tbl.rows[i]._tr.trPr.trHeight is not None
    ]
    assert '54' in heights
    assert '52' in heights


def test_fase_sectie_gebruikt_theme_fontgroottes(monkeypatch) -> None:
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie
    from ui import table_styles

    monkeypatch.setattr(table_styles, 'WORD_BODY_TEXT_SIZE', 12)
    monkeypatch.setattr(table_styles, 'WORD_TABLE_TEXT_SIZE', 6)
    monkeypatch.setattr(table_styles, 'WORD_TABLE_HEADER_SIZE', 9)

    kaart = FaseCard(fase_num=1, stage_name='Fase 1')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    sec = FaseInvoerSectie(id='fase_1', title='Fase 1', fase_card=kaart)

    doc = _export([sec], metadata=ReportMetadata())
    tbl = doc.tables[0]

    assert doc.styles['Normal'].font.size.pt == 12
    assert tbl.rows[1].cells[0].paragraphs[0].runs[0].font.size.pt == 9
    assert tbl.rows[2].cells[0].paragraphs[0].runs[0].font.size.pt == 6


def test_resultaatbeschrijving_specificaties_tabel_in_word() -> None:
    from docx.oxml.ns import qn
    from parsers.models import (
        FileBundle,
        Project,
        ResultPoint,
        ResultStage,
        ResultStep,
        ResultSummary,
        SheetPilingElement,
    )

    project = Project(
        base_name='t',
        project_name='T',
        file_bundle=FileBundle(),
        sheet_piling=[
            SheetPilingElement(
                name='AZ 18 (S240GP)',
                x=0.0,
                bottom=-12.0,
                top=1.0,
                width=1.0,
                opneembaar_moment_knm=250.0,
                steel_quality='S240GP',
            )
        ],
        result_summaries=[
            ResultSummary(
                stage_number=1,
                max_moment_knm=100.0,
                max_shear_kn=50.0,
                max_disp_mm=10.0,
                mob_moment_pct=60.0,
                mob_grond_pct=70.0,
            )
        ],
        result_steps={
            'CUR 166 6.4': ResultStep(
                raw_step='CUR 166 6.4',
                stages={
                    1: ResultStage(
                        stage_number=1,
                        points=[ResultPoint(depth=-2.0, moment=-123.0, shear=45.0, disp=4.0)],
                    )
                },
            ),
            'CUR 166 6.5': ResultStep(
                raw_step='CUR 166 6.5',
                stages={
                    1: ResultStage(
                        stage_number=1,
                        points=[ResultPoint(depth=-1.0, moment=10.0, shear=5.0, disp=-9.0)],
                    )
                },
            ),
        },
    )
    sec = ReportSection(id='conclusietabel', title='Resultaten per fase')

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    fout = WordHoofdstukExporter().export(
        sections=[sec],
        metadata=ReportMetadata(project_name='Test'),
        project=project,
        template_path=None,
        output_path=pad,
    )
    assert fout is None
    doc = Document(pad)
    os.unlink(pad)

    tbl = doc.tables[0]
    assert [col.get(qn('w:w')) for col in tbl._tbl.tblGrid.gridCol_lst] == [
        '2835', '1701', '1134',
    ]
    assert tbl.rows[0].cells[0].text == 'Damwand'
    assert tbl.rows[0].cells[0]._tc.tcPr.tcW.w == 5670
    assert tbl.rows[0].cells[0]._tc.tcPr.find(qn('w:shd')).get(qn('w:fill')) == '147ACF'
    assert [cell.text for cell in tbl.rows[1].cells] == ['Profiel', 'AZ 18', '[-]']
    assert tbl.rows[7].cells[0].text == 'Resultaten'
    assert tbl.rows[7].cells[0]._tc.tcPr.find(qn('w:shd')).get(qn('w:fill')) == '147ACF'
    assert [cell.text for cell in tbl.rows[8].cells] == [
        'Moment Msd', '123,0', '[kNm/m]',
    ]
    assert [cell.text for cell in tbl.rows[10].cells] == [
        'Gemobiliseerd Moment', '60,0', '[%]',
    ]


def test_fase_sectie_voegt_geen_paddingrij_toe_zonder_afbeelding() -> None:
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    kaart = FaseCard(fase_num=1, stage_name='Fase 1')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    sec = FaseInvoerSectie(id='fase_1', title='Fase 1', fase_card=kaart)

    doc = _export([sec], metadata=ReportMetadata())
    tbl = doc.tables[0]
    assert len(tbl.rows) == 3  # twee koprijen + één gevulde tekstregel


def test_fase_sectie_voegt_paddingrij_toe_als_afbeelding_hoger_is(monkeypatch) -> None:
    import base64
    from exporters import word_hoofdstuk_exporter
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    png_1x1 = base64.b64decode(
        'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwAD'
        'hgGAWjR9awAAAABJRU5ErkJggg=='
    )

    monkeypatch.setattr(
        word_hoofdstuk_exporter,
        'render_figuur',
        lambda _img, _project: png_1x1,
    )
    monkeypatch.setattr(word_hoofdstuk_exporter, '_png_hoogte_cm', lambda *_: 12.0)
    kaart = FaseCard(fase_num=1, stage_name='Fase 1')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    sec = FaseInvoerSectie(id='fase_1', title='Fase 1', fase_card=kaart)

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    fout = WordHoofdstukExporter().export(
        sections=[sec],
        metadata=ReportMetadata(),
        project=object(),
        template_path=None,
        output_path=pad,
    )
    assert fout is None
    doc = Document(pad)
    os.unlink(pad)
    tbl = doc.tables[0]
    assert len(tbl.rows) == 4  # twee koprijen + één tekstregel + één paddingrij


def test_fase_sectie_voegt_geen_paddingrij_toe_als_afbeelding_lager_is(monkeypatch) -> None:
    import base64
    from exporters import word_hoofdstuk_exporter
    from reporting.builders.input_description_builder import FaseCard, FaseRow
    from reporting.models import FaseInvoerSectie

    png_1x1 = base64.b64decode(
        'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwAD'
        'hgGAWjR9awAAAABJRU5ErkJggg=='
    )
    monkeypatch.setattr(
        word_hoofdstuk_exporter,
        'render_figuur',
        lambda _img, _project: png_1x1,
    )
    monkeypatch.setattr(word_hoofdstuk_exporter, '_png_hoogte_cm', lambda *_: 0.1)

    kaart = FaseCard(fase_num=1, stage_name='Fase 1')
    kaart.rows.append(FaseRow('Maaiveld Links', '0,9 [m NAP]'))
    sec = FaseInvoerSectie(id='fase_1', title='Fase 1', fase_card=kaart)

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        pad = f.name
    fout = WordHoofdstukExporter().export(
        sections=[sec],
        metadata=ReportMetadata(),
        project=object(),
        template_path=None,
        output_path=pad,
    )
    assert fout is None
    doc = Document(pad)
    os.unlink(pad)
    tbl = doc.tables[0]
    assert len(tbl.rows) == 3  # twee koprijen + één tekstregel, geen paddingrij


# ---------------------------------------------------------------------------
# Taak 6: figuurrendering
# ---------------------------------------------------------------------------

from parsers.models import Project, FileBundle, SheetPilingElement, Stage, Surface


def _mini_project() -> Project:
    return Project(
        base_name='t', project_name='T', file_bundle=FileBundle(),
        sheet_piling=[SheetPilingElement(
            name='AZ 14-700', x=0.0, bottom=-10.0, top=-2.0, width=1.4,
        )],
        stages=[Stage(name='F1')],
        surfaces=[Surface(nr=1, name='MV', points=[{'x': -10, 'y': 0}, {'x': 10, 'y': 0}])],
    )


def test_figuur_placeholder_zonder_project() -> None:
    sec = ReportSection(id='test', title='Grafieken')
    sec.images.append(ReportImageRequest(
        id='fig1', caption='Dwarsdoorsnede fase 1',
        figure_key='section', stage_index=0, step_key=None,
    ))
    doc = _export([sec], metadata=ReportMetadata())
    tekst = ' '.join(p.text for p in doc.paragraphs)
    assert 'Figuur' in tekst


def test_render_figuur_section_geeft_bytes() -> None:
    from reporting.models import ReportImageRequest as RIR
    from exporters.word_hoofdstuk_exporter import WordHoofdstukExporter
    project = _mini_project()
    img_req = RIR(id='f', caption='c', figure_key='section', stage_index=0, step_key=None)
    exp = WordHoofdstukExporter()
    result = exp._render_figuur(img_req, project)
    assert result is not None
    assert result[:4] == b'\x89PNG'


def test_render_figuur_moment_shear_geeft_bytes() -> None:
    from reporting.models import ReportImageRequest as RIR
    from exporters.word_hoofdstuk_exporter import WordHoofdstukExporter
    project = _mini_project()
    img_req = RIR(id='f', caption='c', figure_key='moment_shear', stage_index=0, step_key=None)
    exp = WordHoofdstukExporter()
    result = exp._render_figuur(img_req, project)
    assert result is not None
    assert result[:4] == b'\x89PNG'
